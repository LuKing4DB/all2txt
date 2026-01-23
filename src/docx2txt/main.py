"""
DOCX转TXT脚本（简化版）
将DOCX文件转换为TXT格式，每段一行，图片和表格单独标记
"""

import argparse
import sys
import time
import tempfile
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

# 优先使用相对导入（推荐，作为包安装时使用）
try:
    from .lib.image_detector import has_image
    from .lib.paragraph_processor import process_paragraph
    from .lib.table_processor import process_table
    from .lib.del_docx_auto_num import delete_auto_numbering_in_docx
    from .lib.bookmark_processor import add_bookmarks_to_docx, remove_all_bookmarks
    from ..utils.logger import get_logger
except ImportError:
    # 如果相对导入失败，尝试包绝对导入（作为第三方依赖安装时）
    try:
        from all2txt.docx2txt.lib.image_detector import has_image
        from all2txt.docx2txt.lib.paragraph_processor import process_paragraph
        from all2txt.docx2txt.lib.table_processor import process_table
        from all2txt.docx2txt.lib.del_docx_auto_num import delete_auto_numbering_in_docx
        from all2txt.docx2txt.lib.bookmark_processor import add_bookmarks_to_docx, remove_all_bookmarks
        from all2txt.utils.logger import get_logger
    except ImportError:
        # 如果包导入也失败，使用路径导入（直接运行脚本时）
        # 添加src目录到路径，以便导入lib模块
        sys.path.insert(0, str(Path(__file__).parent.parent))
        sys.path.insert(0, str(Path(__file__).parent))
        from lib.image_detector import has_image
        from lib.paragraph_processor import process_paragraph
        from lib.table_processor import process_table
        from lib.del_docx_auto_num import delete_auto_numbering_in_docx
        from lib.bookmark_processor import add_bookmarks_to_docx, remove_all_bookmarks
        from utils.logger import get_logger

logger = get_logger(__name__)


def has_auto_numbering(docx_path: str) -> bool:
    """
    检测DOCX文件是否包含自动编号
    
    Args:
        docx_path: DOCX文件路径
        
    Returns:
        bool: 如果文档包含自动编号返回True，否则返回False
    """
    try:
        doc = Document(docx_path)
        
        # 方法1: 检查文档中是否有段落使用了 numPr（编号属性）
        for para in doc.paragraphs:
            if para._element.pPr is not None:
                numPr = para._element.pPr.find(qn('w:numPr'))
                if numPr is not None:
                    # 检查 numId 是否不为 0（0 表示无编号）
                    numId_elem = numPr.find(qn('w:numId'))
                    if numId_elem is not None:
                        numId_val = numId_elem.get(qn('w:val'))
                        if numId_val and numId_val != '0':
                            return True
        
        # 方法2: 检查表格中的段落
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para._element.pPr is not None:
                            numPr = para._element.pPr.find(qn('w:numPr'))
                            if numPr is not None:
                                numId_elem = numPr.find(qn('w:numId'))
                                if numId_elem is not None:
                                    numId_val = numId_elem.get(qn('w:val'))
                                    if numId_val and numId_val != '0':
                                        return True
        
        # 方法3: 检查样式中的编号属性
        try:
            from lib.del_docx_auto_num import get_numbering_style_list
            para_id2style = get_numbering_style_list(doc)
            if para_id2style:
                return True
        except:
            pass
        
        return False
    except Exception as e:
        # 如果检测过程中出错，默认返回False，继续正常处理
        logger.warning(f"检测自动编号时出错: {e}")
        return False


def docx_to_txt_simple(docx_path: str, output_path: str = None, debug: bool = False) -> None:
    """
    将DOCX文件转换为TXT文件，每段一行，图片和表格单独标记
    
    Args:
        docx_path: DOCX文件路径
        output_path: 输出TXT文件路径，如果为None则自动生成
        debug: 是否启用debug模式，默认False
    """
    docx_file = Path(docx_path)
    
    if not docx_file.exists():
        logger.error(f"文件不存在: {docx_path}")
        sys.exit(1)
    
    if not docx_file.suffix.lower() == '.docx':
        logger.error(f"不是DOCX文件: {docx_path}")
        sys.exit(1)
    
    # 创建输出文件夹（基于原文件名）
    if output_path is None:
        # 如果没有指定输出路径，使用原文件名创建文件夹
        output_dir = docx_file.parent / docx_file.stem
        output_path = output_dir / (docx_file.stem + '.txt')
        output_type = "自动生成"
        user_specified_output = None
    else:
        user_specified_output = str(output_path)  # 保存用户指定的原始路径
        output_path = Path(output_path)
        # 如果指定了输出路径，创建对应的文件夹
        if output_path.is_dir() or not output_path.suffix:
            # 如果是指定目录或没有扩展名，使用该目录
            output_dir = output_path if output_path.is_dir() else output_path
            output_dir.mkdir(exist_ok=True)
            output_path = output_dir / (docx_file.stem + '.txt')
            output_type = "指定目录"
        else:
            # 如果指定了文件路径，创建同名文件夹
            output_dir = output_path.parent / output_path.stem
            output_path = output_dir / output_path.name
            output_type = "指定文件"
    
    # 确保输出文件夹存在
    output_dir = output_path.parent
    output_dir.mkdir(parents=True, exist_ok=True)
    
    try:
        # 记录开始时间
        start_time = time.time()
        
        # 检测并处理自动编号
        temp_file = None
        actual_docx_path = docx_path
        
        if has_auto_numbering(docx_path):
            logger.info("检测到文档包含自动编号，正在处理...")
            try:
                # 创建临时文件
                temp_file = tempfile.NamedTemporaryFile(
                    mode='wb',
                    suffix='.docx',
                    delete=False,
                    dir=docx_file.parent
                )
                temp_file.close()
                temp_file_path = temp_file.name
                
                # 处理自动编号
                processed_doc = delete_auto_numbering_in_docx(docx_path)
                processed_doc.save(temp_file_path)
                
                actual_docx_path = temp_file_path
                logger.info("✓ 自动编号已处理，生成临时文件继续处理")
            except Exception as e:
                logger.warning(f"处理自动编号时出错: {e}，将使用原始文件继续处理")
                actual_docx_path = docx_path
        
        # 读取DOCX文件（可能是处理后的临时文件）
        doc = Document(actual_docx_path)
        
        # 重要：先删除文档中所有已存在的书签
        # 这确保索引生成和书签生成使用完全相同的element_index编号
        # 因为body级别的书签元素会影响enumerate(doc.element.body)的索引
        removed_bookmark_count = remove_all_bookmarks(doc)
        if removed_bookmark_count > 0:
            logger.debug(f"已从文档中移除 {removed_bookmark_count} 个旧书签元素")
        
        # 添加书签并生成副本（传入同一个doc对象，确保使用相同的文档结构）
        try:
            bookmarked_docx_path = output_dir / (docx_file.stem + '_bookmarked.docx')
            logger.info("正在为文档添加书签...")
            add_bookmarks_to_docx(str(actual_docx_path), str(bookmarked_docx_path), doc=doc)
            logger.info(f"✓ 已生成带书签的文档副本: {bookmarked_docx_path.name}")
        except Exception as e:
            logger.warning(f"添加书签时出错: {e}，将继续处理文档")
        
        # 优化：一次性建立元素到Paragraph对象的映射
        # 这将O(N²)复杂度降低到O(N)，大幅提升性能
        element_to_para = {para._element: para for para in doc.paragraphs}
        
        # 存储提取的内容项（包含索引）
        content_items = []
        # 存储页码信息（索引|页码）
        page_numbers = []
        
        # 统计信息（边处理边统计，避免最后遍历）
        text_count = 0
        image_count = 0
        table_count = 0
        filtered_stamp_count = 0
        page_number_count = 0
        
        # 遍历文档body中的所有元素（按阅读顺序），使用真实的元素位置索引
        # 使用enumerate获取元素在doc.element.body中的真实位置，这样可以通过doc.element.body[index]直接访问
        for element_index, element in enumerate(doc.element.body):
            try:
                # 安全获取tag，处理tag可能是函数或其他类型的情况
                tag = element.tag
                # 如果tag是函数，尝试调用它；如果是字符串，直接使用
                if callable(tag):
                    tag = tag()
                if not isinstance(tag, str):
                    # 如果tag不是字符串，尝试转换为字符串
                    tag = str(tag) if tag is not None else ''
                
                # 处理段落 (w:p)
                if tag.endswith('}p') or tag == 'p':
                    # 优化：先检查是否有图片，避免在process_paragraph内部重复检查
                    element_has_image = has_image(element)
                    items = process_paragraph(doc, element, element_to_para, filter_stamps=True, element_index=element_index, debug=debug)
                    
                    # 优化：边处理边统计，避免最后遍历
                    for item_type, content, idx in items:
                        if item_type == 'page_number':
                            # 页码单独收集
                            page_numbers.append((idx, content))
                            page_number_count += 1
                        else:
                            # 其他内容添加到正常内容列表
                            content_items.append((item_type, content, idx))
                            if item_type == 'text':
                                text_count += 1
                            elif item_type == 'image':
                                image_count += 1
                    
                    # 统计过滤的图片数量（包括印章和无法读取的图片）
                    if element_has_image:
                        # 如果段落有图片但没有输出图片标记，说明被过滤了（可能是印章或无法读取）
                        has_image_marker = any(item_type == 'image' for item_type, _, _ in items)
                        if not has_image_marker:
                            filtered_stamp_count += 1
                
                # 处理表格 (w:tbl)
                elif tag.endswith('}tbl') or tag == 'tbl':
                    items = process_table(element, element_index=element_index)
                    # 优化：边处理边统计
                    for item_type, content, idx in items:
                        content_items.append((item_type, content, idx))
                        table_count += 1
                
                # 其他元素类型（如sectPr等）被跳过，不输出，但索引仍然对应原始位置
            except Exception as e:
                # 如果处理某个元素时出错，记录警告并跳过该元素，继续处理其他元素
                logger.warning(f"处理元素 {element_index} 时出错: {e}，已跳过该元素")
                continue
        
        # 生成索引文件路径和页码文件路径（都在输出文件夹中）
        index_path = output_dir / (output_path.stem + '_index.txt')
        page_number_path = output_dir / (output_path.stem + '_page.txt')
        
        # 统计字数
        total_chars = 0  # 总字符数（去除空格）
        chinese_chars = 0  # 中文字符数
        english_chars = 0  # 英文字符数
        digit_chars = 0  # 数字字符数
        other_chars = 0  # 其他字符数
        
        # 同时写入内容文件和索引文件，并统计字数
        with open(output_path, 'w', encoding='utf-8') as f_content, \
             open(index_path, 'w', encoding='utf-8') as f_index:
            for item_type, content, index in content_items:
                f_content.write(content + '\n')
                f_index.write(str(index) + '\n')
                
                # 只统计文本内容的字数
                if item_type == 'text':
                    # 去除空格后的内容
                    content_no_space = content.replace(' ', '').replace('\t', '')
                    total_chars += len(content_no_space)
                    
                    # 统计各类字符
                    for char in content_no_space:
                        if '\u4e00' <= char <= '\u9fff':  # 中文字符
                            chinese_chars += 1
                        elif char.isalpha() and ord(char) < 128:  # 英文字母
                            english_chars += 1
                        elif char.isdigit():  # 数字
                            digit_chars += 1
                        else:  # 其他字符
                            other_chars += 1
        
        # 写入页码文件（格式：索引序号|页码）
        with open(page_number_path, 'w', encoding='utf-8') as f_page:
            for index, page_num in page_numbers:
                f_page.write(f"{index}|{page_num}\n")
        
        # 记录结束时间并计算耗时
        end_time = time.time()
        elapsed_time = end_time - start_time
        
        # 计算总内容项数（包括页码）
        total_items = len(content_items) + page_number_count
        
        logger.info(f"\n{'='*60}")
        logger.info(f"转换完成")
        logger.info(f"{'='*60}")
        logger.info(f"输入文件: {docx_path}")
        if user_specified_output:
            logger.info(f"输出目录: {user_specified_output} ({output_type})")
        else:
            logger.info(f"输出目录: {output_dir} ({output_type})")
        logger.info(f"\n生成的文件:")
        logger.info(f"  ✓ {output_path.name}          (文本内容)")
        logger.info(f"  ✓ {index_path.name}    (索引信息)")
        if page_number_count > 0:
            logger.info(f"  ✓ {page_number_path.name}     (页码信息)")
        # 检查是否生成了带书签的文档副本
        bookmarked_docx_path = output_dir / (docx_file.stem + '_bookmarked.docx')
        if bookmarked_docx_path.exists():
            logger.info(f"  ✓ {bookmarked_docx_path.name}  (带书签的DOCX副本)")
        logger.info(f"\n完整路径:")
        logger.info(f"  {output_path}")
        logger.info(f"  {index_path}")
        if page_number_count > 0:
            logger.info(f"  {page_number_path}")
        if bookmarked_docx_path.exists():
            logger.info(f"  {bookmarked_docx_path}")
        
        logger.info(f"\n内容统计:")
        logger.info(f"  总计: {total_items} 项")
        logger.info(f"    - 文本段落: {text_count} 个")
        if filtered_stamp_count > 0:
            logger.info(f"    - 图片: {image_count} 个（已过滤 {filtered_stamp_count} 个印章/无效图片）")
        else:
            logger.info(f"    - 图片: {image_count} 个")
        logger.info(f"    - 表格: {table_count} 个")
        if page_number_count > 0:
            logger.info(f"    - 页码: {page_number_count} 个")
        
        logger.info(f"\n字数统计:")
        logger.info(f"  总字符数: {total_chars:,} 个（去除空格）")
        logger.info(f"    - 中文字符: {chinese_chars:,} 个")
        logger.info(f"    - 英文字符: {english_chars:,} 个")
        logger.info(f"    - 数字字符: {digit_chars:,} 个")
        logger.info(f"    - 其他字符: {other_chars:,} 个")
        
        logger.info(f"\n处理耗时: {elapsed_time:.2f} 秒")
        logger.info(f"{'='*60}\n")
        
        # 清理临时文件（如果存在）
        if temp_file is not None:
            try:
                import os
                os.unlink(temp_file.name)
                if debug:
                    logger.debug(f"已清理临时文件: {temp_file.name}")
            except Exception as e:
                if debug:
                    logger.debug(f"清理临时文件时出错: {e}")
        
    except Exception as e:
        logger.error(f"处理文件时出错: {e}")
        import traceback
        traceback.print_exc()
        
        # 确保在出错时也清理临时文件
        if temp_file is not None:
            try:
                import os
                os.unlink(temp_file.name)
            except:
                pass
        
        sys.exit(1)


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description='将DOCX文件转换为TXT文件，每段一行，图片和表格单独标记'
    )
    parser.add_argument(
        'docx_path',
        type=str,
        help='输入的DOCX文件路径'
    )
    parser.add_argument(
        '-o', '--output',
        type=str,
        default=None,
        help='输出目录路径（可选，默认在输入文件同目录下创建同名文件夹）'
    )
    parser.add_argument(
        '--debug',
        action='store_true',
        help='启用debug模式，输出图片处理的详细信息'
    )
    
    args = parser.parse_args()
    docx_to_txt_simple(args.docx_path, args.output, debug=args.debug)


if __name__ == '__main__':
    main()

