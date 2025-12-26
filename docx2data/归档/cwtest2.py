# -*- coding: utf-8 -*-
"""
检测docx文件中哪些段落包含图片
参考CWDocxFile中图片处理方式
"""
import re
import os
from docx import Document


def check_paragraph_has_image(para):
    """
    检查段落是否包含图片
    参考CWRun.init_run_by_run中的图片检测逻辑
    :param para: docx.paragraph对象
    :return: (bool, list) 是否包含图片，以及图片的rId列表
    """
    has_image = False
    pic_rids = []
    
    # 遍历段落中的所有runs
    for run in para.runs:
        try:
            # 标准的python-docx Run对象通过_element.xml访问XML
            # 如果run有xml属性（自定义对象），直接使用；否则使用_element.xml
            if hasattr(run, 'xml'):
                soup_text = str(run.xml)
            elif hasattr(run, '_element') and hasattr(run._element, 'xml'):
                soup_text = str(run._element.xml)
            else:
                continue
            
            # 检查是否有drawing_lst属性（某些自定义对象可能有）
            has_drawing_lst = False
            if hasattr(run, 'drawing_lst') and len(run.drawing_lst) > 0:
                has_drawing_lst = True
            
            # 通过正则表达式检查XML中是否包含drawing标签
            drawing_soup = re.findall(r'<w:drawing>(.*?)</w:drawing>', soup_text, flags=re.DOTALL)
            
            if has_drawing_lst or len(drawing_soup) > 0:
                has_image = True
                # 提取图片的rId
                pattern = r'<a:blip r:embed="(rId\d+)"'
                matches = re.findall(pattern, soup_text)
                pic_rids.extend(matches)
        except Exception as e:
            # 如果处理某个run时出错，继续处理下一个
            continue
    
    return has_image, pic_rids


def detect_images_in_docx(docx_path):
    """
    检测docx文件中哪些段落包含图片
    :param docx_path: docx文件路径
    """
    docx = Document(docx_path)

    print(f"正在检测文档: {docx_path}")
    print("=" * 60)
    
    paragraphs_with_images = []
    
    # 遍历所有段落
    for para_idx, para in enumerate(docx.paragraphs):
        has_image, pic_rids = check_paragraph_has_image(para)
        
        if has_image:
            paragraphs_with_images.append({
                'index': para_idx,
                'text': para.text[:50] + "..." if len(para.text) > 50 else para.text,  # 只显示前50个字符
                'pic_rids': pic_rids
            })
    
    # 输出结果
    if paragraphs_with_images:
        print(f"找到 {len(paragraphs_with_images)} 个包含图片的段落：\n")
        for item in paragraphs_with_images:
            print(f"段落索引: {item['index']}")
            print(f"段落文本预览: {item['text']}")
            print(f"图片rId: {item['pic_rids']}")
            print("-" * 60)
    else:
        print("未找到包含图片的段落")
    
    # 检查表格中的图片
    print("\n检查表格中的图片...")
    table_count = 0
    for table_idx, table in enumerate(docx.tables):
        table_count += 1
        table_has_image = False
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    has_image, pic_rids = check_paragraph_has_image(para)
                    if has_image:
                        table_has_image = True
                        print(f"表格 {table_idx + 1} 的单元格中包含图片，rId: {pic_rids}")
                        break
                if table_has_image:
                    break
            if table_has_image:
                break
    
    if table_count == 0:
        print("文档中没有表格")
    elif not any(check_paragraph_has_image(para)[0] for table in docx.tables for row in table.rows for cell in row.cells for para in cell.paragraphs):
        print("表格中没有找到图片")


if __name__ == "__main__":
    # 在这里设置docx文件路径
    docx_path = r"data/1.docx"  # 修改为你的docx文件路径
    
    # 如果路径是相对路径，可以基于脚本所在目录计算绝对路径
    # script_dir = os.path.dirname(os.path.abspath(__file__))
    # docx_path = os.path.join(script_dir, "..", "example.docx")
    # docx_path = os.path.normpath(docx_path)
    
    detect_images_in_docx(docx_path)
