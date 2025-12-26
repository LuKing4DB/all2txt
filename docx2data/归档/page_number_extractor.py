"""
页码提取模块
从DOCX文档的页脚中提取页码信息，并计算每个段落的页码
"""

import re
from docx import Document
from docx.oxml.ns import qn


def extract_page_numbers_from_footers(doc: Document, debug: bool = False) -> dict:
    """
    从文档页脚中提取页码信息，返回节索引到页码的映射
    
    Args:
        doc: Document对象
        debug: 是否输出调试信息
        
    Returns:
        字典，key为节索引，value为页码（整数），如果读不到则为None
    """
    section_page_map = {}
    
    try:
        # 遍历所有节（section）
        for section_idx, section in enumerate(doc.sections):
            page_num = None
            
            # 检查默认页脚
            try:
                if section.footer is not None:
                    # 遍历页脚中的所有段落
                    for para in section.footer.paragraphs:
                        text = para.text.strip()
                        if text:
                            # 尝试从文本中提取数字（页码）
                            numbers = re.findall(r'^\d+$', text)  # 只匹配纯数字
                            if numbers:
                                try:
                                    page_num = int(numbers[0])
                                    if page_num > 0:
                                        section_page_map[section_idx] = page_num
                                        if debug:
                                            print(f"节 {section_idx + 1}: 从页脚提取到页码 {page_num}")
                                        break
                                except ValueError:
                                    pass
            except Exception as e:
                if debug:
                    print(f"节 {section_idx + 1}: 读取页脚时出错: {e}")
                continue
        
        return section_page_map
        
    except Exception as e:
        # 如果读取页脚时出错，打印错误信息并返回空字典
        import traceback
        print(f"读取页脚时出错: {e}")
        if debug:
            traceback.print_exc()
        return {}


def get_page_breaks(doc: Document) -> list:
    """
    获取文档中所有分页符的位置（段落索引）
    
    Args:
        doc: Document对象
        
    Returns:
        包含分页符位置的段落索引列表
    """
    page_breaks = []
    
    for idx, para in enumerate(doc.paragraphs):
        para_xml = para._element.xml
        
        # 检查段落XML中是否包含分页符 (w:br w:type="page")
        if 'w:type="page"' in para_xml or 'w:type=\'page\'' in para_xml:
            page_breaks.append(idx)
        
        # 检查段落属性中是否有分页前属性
        pPr = para._element.pPr
        if pPr is not None:
            # 检查是否有分页前属性
            pageBreakBefore = pPr.find(qn('w:pageBreakBefore'))
            if pageBreakBefore is not None:
                page_breaks.append(idx)
            
            # 检查是否有分页后属性
            pageBreakAfter = pPr.find(qn('w:pageBreakAfter'))
            if pageBreakAfter is not None:
                page_breaks.append(idx + 1)  # 分页在段落之后
        
        # 检查段落内的运行(run)中是否有分页符
        for run in para.runs:
            run_xml = run._element.xml
            if 'w:type="page"' in run_xml or 'w:type=\'page\'' in run_xml:
                page_breaks.append(idx)
    
    return sorted(set(page_breaks))  # 去重并排序


def get_section_breaks(doc: Document) -> list:
    """
    获取所有分节符的位置（段落索引）
    
    Args:
        doc: Document对象
        
    Returns:
        包含分节符位置的段落索引列表
    """
    section_breaks = []
    
    for idx, para in enumerate(doc.paragraphs):
        para_xml = para._element.xml
        # 检查段落XML中是否包含分节符 (w:sectPr)
        if 'w:sectPr' in para_xml:
            section_breaks.append(idx)
    
    return sorted(set(section_breaks))


def get_page_numbers_for_paragraphs(doc: Document, paragraph_indices: list, debug: bool = False) -> list:
    """
    为每个段落计算页码
    
    Args:
        doc: Document对象
        paragraph_indices: 非空段落在原文档中的索引列表
        
    Returns:
        页码列表，每个段落对应一个页码，如果读不到则为0
    """
    # 从页脚读取每个节的页码
    section_page_map = extract_page_numbers_from_footers(doc, debug=debug)
    
    # 如果没有找到任何页码，所有段落都返回0
    if not section_page_map:
        if debug:
            print("未找到任何页码信息，所有段落返回0")
        return [0] * len(paragraph_indices)
    
    if debug:
        print(f"找到 {len(section_page_map)} 个节的页码信息")
        print(f"节索引范围: {min(section_page_map.keys())} - {max(section_page_map.keys())}")
    
    # 获取所有分节符位置
    section_breaks = get_section_breaks(doc)
    
    if debug:
        print(f"找到 {len(section_breaks)} 个分节符")
    
    # 创建段落索引到页码的映射
    page_numbers = []
    para_idx_in_list = 0
    
    # 遍历所有段落（包括空段落）来确定页码
    for para_idx_in_doc, para in enumerate(doc.paragraphs):
        # 确定当前段落属于哪个节
        # 段落属于它之前最后一个分节符之后的节
        section_idx = 0
        for i, break_idx in enumerate(section_breaks):
            if para_idx_in_doc >= break_idx:
                section_idx = i + 1
            else:
                break
        
        # 获取该节的页码
        if section_idx in section_page_map:
            current_page = section_page_map[section_idx]
        else:
            # 如果当前节没有页码，尝试使用前一个有页码的节
            current_page = 0
            for prev_section_idx in range(section_idx, -1, -1):
                if prev_section_idx in section_page_map:
                    current_page = section_page_map[prev_section_idx]
                    break
        
        # 如果当前段落是非空段落，记录其页码
        if para.text.strip():
            if para_idx_in_list < len(paragraph_indices) and paragraph_indices[para_idx_in_list] == para_idx_in_doc:
                page_numbers.append(current_page if current_page > 0 else 0)
                para_idx_in_list += 1
    
    return page_numbers

