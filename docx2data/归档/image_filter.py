"""
图片过滤模块
用于检测和过滤DOCX文档中的小型图片（如印章等）
支持两种过滤方式：
1. 基于尺寸的过滤（过滤小于指定尺寸的图片）
2. 基于重复检测的过滤（检测重复出现的相同图片，判定为印章）
"""

import hashlib
import io
from pathlib import Path
from typing import Dict, Set, Optional, Tuple
from docx import Document

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

# EMU (English Metric Units) 转换常量
# 1英寸 = 914400 EMU
# 1厘米 ≈ 360000 EMU
EMU_PER_INCH = 914400
EMU_PER_CM = 360000


def has_image(element):
    """
    检查元素中是否包含图片
    
    Args:
        element: XML元素
        
    Returns:
        如果包含图片返回True，否则返回False
    """
    # 检查是否有drawing元素（包含图片）
    drawings = element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
    if drawings:
        return True
    
    # 也检查其他可能的图片位置
    blip = element.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
    if blip is not None:
        return True
    
    return False


def get_image_size_from_paragraph(doc: Document, element):
    """
    从段落中获取图片尺寸（EMU单位）
    使用python-docx的Paragraph对象来获取图片尺寸
    
    Args:
        doc: Document对象
        element: 段落XML元素
        
    Returns:
        (width, height) 元组，单位为EMU，如果无法获取则返回None
    """
    # 找到对应的Paragraph对象
    para_obj = None
    for para in doc.paragraphs:
        if para._element == element:
            para_obj = para
            break
    
    if para_obj is None:
        return None
    
    # 遍历段落中的所有runs，查找包含图片的run
    for run in para_obj.runs:
        # 检查run中是否有图片
        if run._element.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline') is not None:
            # 尝试从run的XML中获取图片尺寸
            inline = run._element.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
            if inline is not None:
                extent = inline.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                if extent is not None:
                    cx = extent.get('cx')  # 宽度
                    cy = extent.get('cy')  # 高度
                    if cx and cy:
                        try:
                            return (int(cx), int(cy))
                        except ValueError:
                            pass
    
    return None


def get_image_hash_from_paragraph(doc: Document, element, debug: bool = False) -> Optional[str]:
    """
    从段落中提取图片的hash值作为特征
    
    Args:
        doc: Document对象
        element: 段落XML元素
        debug: 是否输出调试信息
        
    Returns:
        图片的MD5 hash值（字符串），如果无法获取则返回None
    """
    # 找到对应的Paragraph对象
    para_obj = None
    for para in doc.paragraphs:
        if para._element == element:
            para_obj = para
            break
    
    if para_obj is None:
        if debug:
            print(f"    [DEBUG] 无法找到对应的Paragraph对象")
        return None
    
    # 遍历段落中的所有runs，查找包含图片的run
    for run_idx, run in enumerate(para_obj.runs):
        # 检查run中是否有图片
        inline = run._element.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
        if inline is not None:
            # 查找图片的relationship ID
            blip = inline.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
            if blip is not None:
                from docx.oxml.ns import qn
                
                # 方法1: 尝试获取嵌入的图片 (r:embed)
                r_embed = blip.get(qn('r:embed'))
                if r_embed:
                    try:
                        # 获取图片数据
                        image_part = doc.part.related_parts[r_embed]
                        image_bytes = image_part.blob
                        # 计算MD5 hash
                        image_hash = hashlib.md5(image_bytes).hexdigest()
                        if debug:
                            print(f"    [DEBUG] 成功获取图片hash (run {run_idx}, r:embed={r_embed})")
                        return image_hash
                    except KeyError as e:
                        if debug:
                            print(f"    [DEBUG] 无法在related_parts中找到图片 (r:embed={r_embed}): {e}")
                    except AttributeError as e:
                        if debug:
                            print(f"    [DEBUG] 图片部分没有blob属性 (r:embed={r_embed}): {e}")
                    except Exception as e:
                        if debug:
                            print(f"    [DEBUG] 获取嵌入图片时出错 (r:embed={r_embed}): {e}")
                
                # 方法2: 检查是否是链接图片 (r:link) - 链接图片无法获取hash
                r_link = blip.get(qn('r:link'))
                if r_link:
                    if debug:
                        print(f"    [DEBUG] 检测到链接图片 (r:link={r_link})，无法获取hash（图片数据在外部）")
                    return None
                
                # 方法3: 尝试通过其他方式获取图片
                if debug:
                    print(f"    [DEBUG] blip元素存在但未找到r:embed或r:link属性")
                    print(f"    [DEBUG] blip属性: {dict(blip.attrib)}")
            else:
                if debug:
                    print(f"    [DEBUG] 未找到blip元素")
        else:
            # 尝试其他可能的图片位置
            # 检查是否有其他类型的图片元素
            anchor = run._element.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
            if anchor is not None:
                if debug:
                    print(f"    [DEBUG] 检测到anchor类型的图片，当前不支持")
    
    if debug:
        print(f"    [DEBUG] 遍历完所有runs，未找到可获取hash的图片")
    return None


def get_image_data_from_paragraph(doc: Document, element, debug: bool = False) -> Optional[Tuple[bytes, str]]:
    """
    从段落中提取图片数据
    
    Args:
        doc: Document对象
        element: 段落XML元素
        debug: 是否输出调试信息
        
    Returns:
        (图片字节数据, 文件扩展名) 元组，如果无法获取则返回None
    """
    # 找到对应的Paragraph对象
    para_obj = None
    for para in doc.paragraphs:
        if para._element == element:
            para_obj = para
            break
    
    if para_obj is None:
        if debug:
            print(f"    [DEBUG] 无法找到对应的Paragraph对象")
        return None
    
    # 遍历段落中的所有runs，查找包含图片的run
    for run_idx, run in enumerate(para_obj.runs):
        # 检查run中是否有图片
        inline = run._element.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
        if inline is not None:
            # 查找图片的relationship ID
            blip = inline.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
            if blip is not None:
                from docx.oxml.ns import qn
                
                # 尝试获取嵌入的图片 (r:embed)
                r_embed = blip.get(qn('r:embed'))
                if r_embed:
                    try:
                        # 获取图片数据
                        image_part = doc.part.related_parts[r_embed]
                        image_bytes = image_part.blob
                        
                        # 获取图片扩展名
                        # 从content_type推断扩展名
                        content_type = image_part.content_type
                        ext_map = {
                            'image/png': '.png',
                            'image/jpeg': '.jpg',
                            'image/jpg': '.jpg',
                            'image/gif': '.gif',
                            'image/bmp': '.bmp',
                            'image/x-emf': '.emf',
                            'image/x-wmf': '.wmf',
                        }
                        ext = ext_map.get(content_type, '.png')  # 默认使用png
                        
                        if debug:
                            print(f"    [DEBUG] 成功获取图片数据 (run {run_idx}, r:embed={r_embed}, type={content_type})")
                        return (image_bytes, ext)
                    except KeyError as e:
                        if debug:
                            print(f"    [DEBUG] 无法在related_parts中找到图片 (r:embed={r_embed}): {e}")
                    except AttributeError as e:
                        if debug:
                            print(f"    [DEBUG] 图片部分没有blob属性 (r:embed={r_embed}): {e}")
                    except Exception as e:
                        if debug:
                            print(f"    [DEBUG] 获取嵌入图片时出错 (r:embed={r_embed}): {e}")
                
                # 链接图片无法获取数据
                r_link = blip.get(qn('r:link'))
                if r_link:
                    if debug:
                        print(f"    [DEBUG] 检测到链接图片 (r:link={r_link})，无法获取数据（图片数据在外部）")
                    return None
    
    if debug:
        print(f"    [DEBUG] 遍历完所有runs，未找到可获取数据的图片")
    return None


def is_square_image(doc: Document, element, aspect_ratio_tolerance: float = 0.3, debug: bool = False) -> bool:
    """
    检查图片是否近似方形（长宽比接近1:1）
    
    Args:
        doc: Document对象
        element: 段落XML元素
        aspect_ratio_tolerance: 长宽比容差，默认0.3（即长宽比在0.7-1.3之间认为是方形）
        debug: 是否输出调试信息
        
    Returns:
        如果图片近似方形返回True，否则返回False
    """
    size = get_image_size_from_paragraph(doc, element)
    if size is None:
        if debug:
            print(f"    [DEBUG] 无法获取图片尺寸，无法判断是否为方形")
        return False
    
    width_emu, height_emu = size
    if width_emu == 0 or height_emu == 0:
        return False
    
    # 计算长宽比（宽/高）
    aspect_ratio = width_emu / height_emu
    
    # 如果长宽比接近1:1（在容差范围内），认为是方形
    is_square = abs(aspect_ratio - 1.0) <= aspect_ratio_tolerance
    
    if debug:
        print(f"    [DEBUG] 图片尺寸: {width_emu}x{height_emu} EMU, 长宽比: {aspect_ratio:.2f}, 是否方形: {is_square}")
    
    return is_square


def is_red_image(doc: Document, element, red_threshold: float = 0.3, debug: bool = False) -> bool:
    """
    检查图片是否主要是红色
    
    Args:
        doc: Document对象
        element: 段落XML元素
        red_threshold: 红色像素占比阈值，默认0.3（30%以上红色像素认为是红色图片）
        debug: 是否输出调试信息
        
    Returns:
        如果图片主要是红色返回True，否则返回False
    """
    if not PIL_AVAILABLE:
        if debug:
            print(f"    [DEBUG] PIL库未安装，无法检测图片颜色")
        return False
    
    # 获取图片数据
    image_data = get_image_data_from_paragraph(doc, element, debug=debug)
    if image_data is None:
        if debug:
            print(f"    [DEBUG] 无法获取图片数据，无法检测颜色")
        return False
    
    image_bytes, ext = image_data
    
    try:
        # 打开图片
        img = Image.open(io.BytesIO(image_bytes))
        
        # 转换为RGB模式（如果不是）
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        # 获取图片尺寸
        width, height = img.size
        total_pixels = width * height
        
        if total_pixels == 0:
            return False
        
        # 统计红色像素数量
        red_pixel_count = 0
        
        # 采样检测（如果图片太大，采样检测以提高性能）
        sample_step = max(1, min(width, height) // 100)  # 采样步长
        
        for y in range(0, height, sample_step):
            for x in range(0, width, sample_step):
                r, g, b = img.getpixel((x, y))
                
                # 判断是否为红色像素
                # 红色特征：R值较高，且R > G和R > B
                if r > 100 and r > g * 1.2 and r > b * 1.2:
                    red_pixel_count += 1
        
        # 计算红色像素占比
        sampled_pixels = ((width // sample_step) * (height // sample_step))
        red_ratio = red_pixel_count / sampled_pixels if sampled_pixels > 0 else 0
        
        is_red = red_ratio >= red_threshold
        
        if debug:
            print(f"    [DEBUG] 图片尺寸: {width}x{height}, 采样像素: {sampled_pixels}, 红色像素: {red_pixel_count}, 红色占比: {red_ratio:.2%}, 是否红色: {is_red}")
        
        return is_red
        
    except Exception as e:
        if debug:
            print(f"    [DEBUG] 检测图片颜色时出错: {e}")
        return False


def has_stamp_features(
    doc: Document, 
    element, 
    check_square: bool = True,
    check_red: bool = True,
    aspect_ratio_tolerance: float = 0.3,
    red_threshold: float = 0.3,
    debug: bool = False
) -> Tuple[bool, Dict[str, bool]]:
    """
    检查图片是否具有印章特征（近似方形或红色）
    
    Args:
        doc: Document对象
        element: 段落XML元素
        check_square: 是否检查是否为方形，默认True
        check_red: 是否检查是否为红色，默认True
        aspect_ratio_tolerance: 长宽比容差，默认0.3
        red_threshold: 红色像素占比阈值，默认0.3
        debug: 是否输出调试信息
        
    Returns:
        (是否具有印章特征, 特征详情字典) 元组
    """
    features = {
        'is_square': False,
        'is_red': False
    }
    
    if check_square:
        features['is_square'] = is_square_image(doc, element, aspect_ratio_tolerance, debug=debug)
    
    if check_red:
        features['is_red'] = is_red_image(doc, element, red_threshold, debug=debug)
    
    # 只要满足任一特征就认为是印章
    has_features = features['is_square'] or features['is_red']
    
    return (has_features, features)


def save_stamp_images(
    doc: Document,
    stamp_hashes: Set[str],
    output_dir: str,
    debug: bool = False
) -> Dict[str, str]:
    """
    保存印章图片到磁盘
    
    Args:
        doc: Document对象
        stamp_hashes: 被判定为印章的图片hash集合
        output_dir: 输出目录路径
        debug: 是否输出调试信息
        
    Returns:
        字典，key为hash值，value为保存的文件路径
    """
    if not stamp_hashes:
        return {}
    
    # 创建输出目录
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    saved_files: Dict[str, str] = {}
    saved_hashes: Set[str] = set()  # 已保存的hash，避免重复保存
    
    # 遍历所有段落，查找印章图片
    for para_idx, para in enumerate(doc.paragraphs):
        if has_image(para._element):
            image_hash = get_image_hash_from_paragraph(doc, para._element, debug=False)
            if image_hash and image_hash in stamp_hashes and image_hash not in saved_hashes:
                # 获取图片数据
                image_data = get_image_data_from_paragraph(doc, para._element, debug=debug)
                if image_data:
                    image_bytes, ext = image_data
                    # 生成文件名
                    filename = f"stamp_{image_hash[:8]}{ext}"
                    filepath = output_path / filename
                    
                    # 保存文件
                    try:
                        with open(filepath, 'wb') as f:
                            f.write(image_bytes)
                        saved_files[image_hash] = str(filepath)
                        saved_hashes.add(image_hash)
                        if debug:
                            print(f"  已保存印章图片: {filepath} (hash: {image_hash[:8]}...)")
                    except Exception as e:
                        if debug:
                            print(f"  保存印章图片失败 (hash: {image_hash[:8]}...): {e}")
    
    return saved_files


def save_non_stamp_image(
    doc: Document,
    element,
    output_dir: str,
    image_index: int = 0,
    debug: bool = False
) -> Optional[str]:
    """
    保存单个非印章图片到磁盘
    
    Args:
        doc: Document对象
        element: 段落XML元素
        output_dir: 输出目录路径
        image_index: 图片序号（用于生成文件名）
        debug: 是否输出调试信息
        
    Returns:
        保存的文件路径（相对路径），如果保存失败返回None
    """
    # 获取图片数据
    image_data = get_image_data_from_paragraph(doc, element, debug=debug)
    if image_data is None:
        if debug:
            print(f"    [DEBUG] 无法获取图片数据，无法保存")
        return None
    
    image_bytes, ext = image_data
    
    # 创建输出目录
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    # 生成文件名（使用序号和hash）
    image_hash = get_image_hash_from_paragraph(doc, element, debug=False)
    if image_hash:
        filename = f"image_{image_index:04d}_{image_hash[:8]}{ext}"
    else:
        filename = f"image_{image_index:04d}{ext}"
    
    filepath = output_path / filename
    
    # 保存文件
    try:
        with open(filepath, 'wb') as f:
            f.write(image_bytes)
        if debug:
            print(f"  已保存图片: {filepath}")
        # 返回相对路径（相对于输出文本文件的路径）
        return str(filepath)
    except Exception as e:
        if debug:
            print(f"  保存图片失败: {e}")
        return None


def detect_stamp_images_by_features(
    doc: Document,
    check_red: bool = True,
    red_threshold: float = 0.3,
    debug: bool = False
) -> Set[str]:
    """
    通过特征检测判定印章图片（仅使用红色特征检测）
    判定条件：大部分像素为红色的图片
    
    Args:
        doc: Document对象
        check_red: 是否检查是否为红色，默认True
        red_threshold: 红色像素占比阈值，默认0.3（30%以上红色像素认为是印章）
        debug: 是否输出调试信息
        
    Returns:
        被判定为印章的图片hash集合（这些hash对应的图片在全文都会被过滤）
    """
    stamp_hashes: Set[str] = set()
    # 存储每个hash对应的段落元素，用于特征检测
    hash_to_elements: Dict[str, list] = {}
    
    # 先找出所有包含图片的段落
    paragraphs_with_images = []
    for para in doc.paragraphs:
        if has_image(para._element):
            paragraphs_with_images.append(para)
    
    total_image_paragraphs = len(paragraphs_with_images)
    print(f"正在扫描全文检测印章图片（共 {total_image_paragraphs} 个包含图片的段落）...")
    print(f"使用特征检测：红色（红色像素占比 >= {red_threshold*100:.0f}%）")
    
    # 收集所有图片的hash和对应的元素
    # 注意：即使无法获取hash的图片，也要检测其是否为印章
    for para in paragraphs_with_images:
        image_hash = get_image_hash_from_paragraph(doc, para._element, debug=debug)
        if image_hash:
            if image_hash not in hash_to_elements:
                hash_to_elements[image_hash] = []
            hash_to_elements[image_hash].append(para._element)
        else:
            # 无法获取hash的图片，使用None作为key，单独处理
            # 这样可以确保所有图片都被检测
            if None not in hash_to_elements:
                hash_to_elements[None] = []
            hash_to_elements[None].append(para._element)
    
    # 检测具有红色特征的图片
    feature_stamp_count = 0
    for image_hash, elements in hash_to_elements.items():
        # 检查第一个元素是否具有红色特征
        if elements:
            has_features, features = has_stamp_features(
                doc, 
                elements[0],
                check_square=False,  # 不使用方形检测
                check_red=check_red,
                red_threshold=red_threshold,
                debug=debug
            )
            
            if has_features:
                # 如果图片有hash，使用hash；如果没有hash，使用特殊标记
                if image_hash:
                    stamp_hashes.add(image_hash)
                    print(f"  检测到印章图片 (hash: {image_hash[:8]}...)，具有红色特征，将在全文过滤")
                else:
                    # 无法获取hash的图片，但检测到红色特征，需要特殊处理
                    # 使用一个特殊标记来标识这些图片
                    # 但由于无法用hash标识，我们需要在后续处理中直接检测
                    if debug:
                        print(f"  [警告] 检测到红色图片但无法获取hash，将在后续处理中直接检测")
                feature_stamp_count += 1
    
    if feature_stamp_count > 0:
        print(f"通过红色特征检测发现 {feature_stamp_count} 个印章图片")
    else:
        print("未检测到具有红色特征的图片")
    
    return stamp_hashes


def is_large_image(doc: Document, element, min_size_cm: float = 2.0) -> bool:
    """
    检查元素中的图片是否足够大（不是印章等小型图片）
    
    Args:
        doc: Document对象
        element: XML元素（段落）
        min_size_cm: 最小尺寸阈值（厘米），默认2厘米
        
    Returns:
        如果包含足够大的图片返回True，否则返回False
    """
    size = get_image_size_from_paragraph(doc, element)
    if size is None:
        # 如果无法获取尺寸，默认保留（可能是其他类型的图片）
        return True
    
    width_emu, height_emu = size
    min_size_emu = min_size_cm * EMU_PER_CM
    
    # 如果宽度或高度小于阈值，认为是小型图片（印章），过滤掉
    if width_emu < min_size_emu or height_emu < min_size_emu:
        return False
    
    return True


def is_not_stamp_image(doc: Document, element, stamp_hashes: Set[str], debug: bool = False) -> Tuple[bool, Optional[str]]:
    """
    检查图片是否不是印章（通过hash值判断）
    如果图片hash在stamp_hashes集合中，说明是印章，返回False（过滤掉）
    
    Args:
        doc: Document对象
        element: XML元素（段落）
        stamp_hashes: 被判定为印章的图片hash集合（采样检测到的重复图片）
        debug: 是否输出调试信息
        
    Returns:
        (是否保留, 图片hash) 元组，如果保留返回(True, hash)，如果过滤返回(False, hash)
    """
    if not stamp_hashes:
        # 如果没有检测到印章，全部保留
        return (True, None)
    
    image_hash = get_image_hash_from_paragraph(doc, element, debug=debug)
    if image_hash is None:
        # 如果无法获取hash，默认保留（避免误过滤）
        if debug:
            print(f"    [DEBUG] 无法获取图片hash，保留")
        return (True, None)
    
    # 如果图片hash在印章集合中，说明是印章，过滤掉
    if image_hash in stamp_hashes:
        if debug:
            print(f"    [DEBUG] 过滤印章图片 (hash: {image_hash[:8]}...)")
        return (False, image_hash)
    
    # 不在印章集合中，保留
    if debug:
        print(f"    [DEBUG] 保留图片 (hash: {image_hash[:8]}...)")
    return (True, image_hash)

