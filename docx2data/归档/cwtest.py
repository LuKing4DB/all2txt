#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
商务标完整测试脚本（整合版）
将所有依赖函数整合到本脚本中，不依赖其他本地模块
"""

import os
import sys
import json
import time
import re
import copy
import shutil
import traceback
from pathlib import Path
from enum import Enum, auto
from typing import Dict, List, Tuple, Any, Optional
from collections import Counter
from concurrent.futures import ThreadPoolExecutor, as_completed
from urllib.parse import urlparse

# 设置标准输出和标准错误为UTF-8编码，解决Windows PowerShell中文乱码问题
if sys.platform == 'win32':
    try:
        if hasattr(sys.stdout, 'reconfigure'):
            sys.stdout.reconfigure(encoding='utf-8')
        if hasattr(sys.stderr, 'reconfigure'):
            sys.stderr.reconfigure(encoding='utf-8')
        os.environ['PYTHONIOENCODING'] = 'utf-8'
    except Exception:
        pass

# 尝试导入第三方库
try:
    import docx
    from docx.document import Document
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
    from docx.shared import Pt, Cm
    from docx.table import Table
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    print("警告: python-docx模块未安装，无法处理DOCX文件")
    DOCX_AVAILABLE = False

try:
    import requests
    REQUESTS_AVAILABLE = True
except ImportError:
    print("警告: requests模块未安装，无法进行文件转换")
    REQUESTS_AVAILABLE = False

try:
    from openai import OpenAI, RateLimitError, APIConnectionError, APIError
    OPENAI_AVAILABLE = True
except ImportError:
    print("警告: openai模块未安装，无法调用大模型")
    OPENAI_AVAILABLE = False


LORA_API_KEY = "sk-6AciH8uCHs7JMW7mycL6vZDmFnhOwsPeO72G1d4F7g2ioko3"
LORA_URL = "http://36.213.210.26:3000/v1"
LORA_MODEL_NAME = "qwen3-4bl-title"
lora_client = OpenAI(api_key=LORA_API_KEY, base_url=LORA_URL)


# =============================================================================
# 文件类型枚举和配置
# =============================================================================

class ZBFileType(Enum):
    PDF = auto()
    DOC = auto()
    DOCX = auto()
    UNKNOWN = auto()


ZB_FILE_TYPE_CONFIG = {".docx": ZBFileType.DOCX, ".pdf": ZBFileType.PDF, ".doc": ZBFileType.DOC}


# =============================================================================
# JSON提取相关函数
# =============================================================================

def get_longest_json_string(string):
    """从最外层 {} 进行切割，并返回最长的 {} 结构"""
    result = ""
    buffer = ""
    count = 0
    for char in string:
        if char == '{':
            count += 1
        elif char == '}':
            count -= 1
            if count == 0 and len(result) < len(buffer):
                result = buffer + '}'
                buffer = ""
        if count > 0:
            buffer += char
    return result


def insert_missing_commas(json_str):
    """使用正则表达式在缺失逗号的位置插入逗号"""
    pattern = r'(":\s*"[^"]*)"\s*(")'
    replacement = r'\1", \2'
    previous_str = None
    while previous_str != json_str:
        previous_str = json_str
        json_str = re.sub(pattern, replacement, json_str)
    return json_str


def fix_string(string):
    """处理未规格化的回车空行"""
    b_single_in_string = False
    b_double_in_string = False
    new_string = ""
    count = 0
    for char in string:
        if char == "{":
            count += 1
        if char == '"' and not b_single_in_string:
            b_double_in_string = not b_double_in_string
        if char == "'" and not b_double_in_string:
            b_single_in_string = not b_single_in_string
        if not b_double_in_string and not b_single_in_string:
            if char == '\n' or char == ' ':
                continue
        if (count > 0 and not b_single_in_string and not b_double_in_string
                and char not in [':', '"', "'", '{', '}', ',', '，', '-', '[', ']']
                and not char.isdigit()):
            if char == '：':
                new_string += ':'
            if char == '，':
                new_string += ','
            continue
        if char == '}':
            if count <= 0:
                continue
            count -= 1
            if b_single_in_string or b_double_in_string:
                new_string += '"'
            b_single_in_string = False
            b_double_in_string = False
        new_string += '"' if char == "'" else char
    if count > 0:
        for _ in range(count):
            new_string += "}"
    return new_string


def pre_fix_string(string):
    """针对引号的修复"""
    buff_string = ""
    new_string = ""
    single = 0
    chars_set = [',', '{', '}', '[', ']']
    after_colon = False
    for char in string:
        if char in chars_set or (not after_colon and char == ':'):
            if single % 2 == 1:
                buff_string = (buff_string.replace('"', '')
                               .replace("'", '')
                               .replace("\n", "")
                               .strip())
                new_string += f'"{buff_string}"' + char
                buff_string = ""
                single = 0
            else:
                if single == 0:
                    buff_string = buff_string.strip()
                    if not buff_string.strip().isdigit() and not buff_string == '':
                        buff_string = f'"{buff_string}"'
                new_string += buff_string + char
                buff_string = ""
                single = 0
        else:
            if char == '"' or char == "'":
                if single:
                    temp = buff_string.split('"')[-1]
                    new_string += f'"{temp}"'
                    buff_string = ""
                else:
                    buff_string += char
                single += 1
            else:
                buff_string += char

        if char == ':':
            after_colon = True
        elif char in chars_set:
            after_colon = False

    if single % 2 == 1:
        buff_string = (buff_string.replace('"', '')
                       .replace("'", '')
                       .replace("\n", "")
                       .strip())
        new_string += f'"{buff_string}"'
        buff_string = ""
    else:
        if single == 0:
            buff_string = buff_string.strip()
            if not buff_string.strip().isdigit() and not buff_string == '':
                buff_string = f'"{buff_string}"'
        new_string += buff_string
        buff_string = ""
    new_string += buff_string
    return new_string


def insert_missing_quotation(new_string):
    """可能会出现引号缺失及括号匹配问题"""
    buff_string = ""
    string = ""
    colon_num = 0
    for char in new_string:
        buff_string += char
        if char == "{" or char == "}":
            if colon_num > 0:
                string += buff_string
                buff_string = ""
                colon_num = 0
            else:
                string += char
                buff_string = ""
                colon_num = 0
        if char == ":":
            colon_num += 1
    string += buff_string
    for i in range(colon_num):
        string += '}'
    return string


def br_to_next_line(fixing_dict: dict):
    """处理所有的<br>改成换行符"""
    keys_to_rename = []
    for key in list(fixing_dict.keys()):
        value = fixing_dict[key]
        if isinstance(value, str):
            fixing_dict[key] = value.replace('<br>', '\n')
        elif isinstance(value, dict):
            br_to_next_line(value)
        if isinstance(key, str):
            fixed_key_string = key.replace('<br>', '\n')
            if fixed_key_string != key:
                keys_to_rename.append((key, fixed_key_string))
    for old_key, new_key in keys_to_rename:
        if old_key in fixing_dict:
            fixing_dict[new_key] = copy.deepcopy(fixing_dict[old_key])
            del fixing_dict[old_key]
    return fixing_dict


def split_fix_dcit(string):
    """拆分修复函数，以逗号 , 分隔"""
    brace_num = -1
    start_index = -1
    buff_string = ""
    string_list = []
    for idx, char in enumerate(string):
        if char == "{" or char == "[":
            brace_num += 1
            if brace_num == 1:
                start_index = idx
        elif char == '}' or char == "]":
            brace_num -= 1
            if brace_num == 0:
                buff_string = dict_string_fix(buff_string + '"<502JGlue>"').replace('"<502JGlue>"', "")
                buff_string += split_fix_dcit(string[start_index:idx+1]) if char == '}' else split_fix_list(string[start_index:idx+1])
        elif char == ',' and brace_num == 0:
            string_list.append(dict_string_fix(buff_string))
            buff_string = ""
        elif brace_num == 0:
            buff_string += char
    if buff_string:
        string_list.append(dict_string_fix(buff_string))
    new_string = [s for s in string_list if not s.strip() == ""]
    return '{' + f'{",".join(new_string)}' + '}'


def split_fix_list(string):
    """针对列表的处理"""
    brace_num = -1
    start_index = -1
    buff_string = ""
    string_list = []
    for idx, char in enumerate(string):
        if char == "{" or char == "[":
            brace_num += 1
            if brace_num == 1:
                start_index = idx
        elif char == '}' or char == "]":
            brace_num -= 1
            if brace_num == 0:
                buff_string += split_fix_dcit(string[start_index:idx+1]) if char == '}' else split_fix_list(string[start_index:idx+1])
        elif char == ',' and brace_num == 0:
            string_list.append(list_string_fix(buff_string))
            buff_string = ""
        elif brace_num == 0:
            buff_string += char
    if buff_string:
        string_list.append(list_string_fix(buff_string))
    new_string = [s for s in string_list if not s.strip() == ""]
    return '[' + f'{",".join(new_string)}' + ']'


def dict_string_fix(splits_slice):
    """修复字符串键与字符串值对"""
    splits_slice = splits_slice.strip()
    try:
        json.loads("{" + splits_slice + "}", strict=False)
        return splits_slice.replace("\\'", "").replace("'", "")
    except json.JSONDecodeError:
        if not splits_slice[0] == '"':
            splits_slice = '"' + splits_slice
        if not splits_slice[-1] == '"':
            splits_slice = splits_slice + '"'
        try:
            json.loads("{" + splits_slice + "}", strict=False)
            return splits_slice.replace("\\'", "").replace("'", "")
        except json.JSONDecodeError:
            last_char = '.'
            colon_idx = -1
            last_char_idx = -1
            for idx, char in enumerate(splits_slice):
                if char == ':' and last_char == '"':
                    colon_idx = idx
                    break
                elif char == '"' and last_char == ':':
                    colon_idx = last_char_idx
                    break
                if not char == ' ':
                    last_char = char
                    last_char_idx = idx
            if colon_idx == -1:
                for idx, char in enumerate(splits_slice):
                    if char == ':':
                        colon_idx = idx
                        break
            if colon_idx == -1:
                for idx, char in enumerate(splits_slice):
                    if char == ' ':
                        colon_idx = idx
                        break
            if colon_idx == -1:
                tmp = splits_slice.replace('"', "")
                splits_slice = f'"{tmp}":"{tmp}"'
            else:
                left = splits_slice[:colon_idx].replace('"', '').strip()
                right = splits_slice[colon_idx + 1:].replace('"', '').strip()
                splits_slice = f'"{left}":"{right}"'
            return splits_slice.replace("\\'", "").replace("'", "")


def list_string_fix(item):
    """修复列表中的元素"""
    if item.strip() == "":
        return item
    item = item.strip()
    pattern = r'^[-+]?[0-9]*\.?[0-9]+([eE][-+]?[0-9]+)?$'
    match = re.match(pattern, item)
    if match:
        return item
    if item[0] == '"' and item[-1] == '"':
        return item
    if item[0] == '{' and item[-1] == '}' or item[0] == '[' and item[-1] == ']':
        return item
    return '"' + item.replace('"', '').replace('\"', '') + '"'


def extract_json_psola(string):
    """提取JSON对象"""
    if not string.strip():
        return {}
    string1 = insert_missing_quotation(string)
    string1 = fix_string(string1)
    match = get_longest_json_string(string1)
    if match:
        try:
            json_data = match.replace('\n', '<br>')
            return br_to_next_line(json.loads(json_data))
        except json.JSONDecodeError:
            try:
                string = pre_fix_string(string)
                string = insert_missing_quotation(string)
                string = fix_string(string)
                match = get_longest_json_string(string)
                json_data = match.replace('\n', '<br>')
                fixed_json = insert_missing_commas(json_data)
                return br_to_next_line(json.loads(fixed_json))
            except json.JSONDecodeError as fixed_error:
                return {"Error": "Fix Error!"}
    else:
        return {"Error": "Match Error!"}


def extract_json_cw(s):
    """使用正则表达式提取JSON"""
    try:
        json_match = re.search(r'\{(?:[^{}]*|(?R))*\}', s, re.DOTALL)
        if json_match:
            json_str = json_match.group()
        else:
            return {}
        return json.loads(json_str, strict=False)
    except json.JSONDecodeError as e:
        print(f"Invalid JSON format: {e}")
        print(s)
        return {}


def extract_json_from_string(string):
    """最终调用，提取JSON字符串"""
    string = string.replace('\\"', '"').replace("::", ":")
    result = extract_json_psola(string)
    if 'Error' in result:
        result = extract_json_psola(split_fix_dcit(string))
        if 'Error' in result:
            return extract_json_cw(string)
    return result


# =============================================================================
# 文档处理相关函数（简化版）
# =============================================================================

def delete_auto_numbering_in_docx(docx_path):
    """
    删除docx文档中所有段落的自动编号并替换为静态编号（简化版）
    注意：这是一个简化版本，完整版本需要更多依赖
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx模块未安装，无法处理DOCX文件")
    
    try:
        doc = docx.Document(docx_path)
        # 简化处理：直接返回文档对象
        # 完整版本需要处理自动编号替换，这里简化处理
        return doc
    except Exception as e:
        print(f"加载文档时出错: {e}")
        raise


def preprocess_docx_breaks(document, save_path=None):
    """预处理DOCX文档，删除分页符（保留分节符）"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx模块未安装，无法处理DOCX文件")
    
    print("开始预处理分页符（保留分节符）...")
    preprocess_start = time.time()

    paragraphs = list(document.paragraphs)
    removed_count = 0

    for paragraph in paragraphs:
        br_elements = paragraph._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br')
        has_page_break = False
        for br in br_elements:
            br_type = br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
            if br_type in ['page', 'nextPage', 'evenPage', 'oddPage']:
                has_page_break = True
                break
        
        if has_page_break:
            _remove_breaks_from_element(paragraph._element)
            removed_count += 1

    # 处理连续空行
    print("开始处理连续空行...")
    paragraphs_list = list(document.paragraphs)
    paragraphs_to_delete = []
    
    i = 0
    while i < len(paragraphs_list):
        if paragraphs_list[i].text.strip() == "":
            empty_start = i
            empty_end = i
            while empty_end + 1 < len(paragraphs_list) and paragraphs_list[empty_end + 1].text.strip() == "":
                empty_end += 1
            
            empty_count = empty_end - empty_start + 1
            if empty_count >= 3:
                for j in range(empty_start + 2, empty_end + 1):
                    paragraphs_to_delete.append(paragraphs_list[j])
            
            i = empty_end + 1
        else:
            i += 1
    
    empty_removed_count = 0
    for paragraph in reversed(paragraphs_to_delete):
        try:
            paragraph._element.getparent().remove(paragraph._element)
            empty_removed_count += 1
        except Exception as e:
            print(f"删除空行段落时出错: {str(e)}")
    
    if empty_removed_count > 0:
        print(f"删除了 {empty_removed_count} 个多余的空行段落")

    preprocess_time = time.time() - preprocess_start
    print(f"预处理完成，耗时: {preprocess_time:.3f}秒")
    print(f"删除了 {removed_count} 个段落中的分页符（分节符已保留）")

    if save_path:
        try:
            print(f"保存预处理后的文档到: {save_path}")
            document.save(save_path)
            print("文档保存成功！")
        except Exception as e:
            print(f"保存文档时出错: {str(e)}")

    return document


def _remove_breaks_from_element(element):
    """从element中移除分页符（保留分节符）"""
    try:
        br_elements = element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br')
        for br in br_elements:
            br_type = br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
            if br_type in ['page', 'nextPage', 'evenPage', 'oddPage']:
                parent = br.getparent()
                if parent is not None:
                    parent.remove(br)
    except Exception:
        pass


# =============================================================================
# 文档处理核心函数
# =============================================================================

def _extract_text_from_element(element):
    """直接从element中提取文本内容"""
    try:
        text_parts = []
        for run in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
            run_text = ""
            for text_elem in run.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                if text_elem.text:
                    run_text += text_elem.text

            br_elements = run.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br')
            if br_elements:
                run_text += "<br>"

            text_parts.append(run_text)

        return ''.join(text_parts)
    except Exception:
        return ""


def _has_page_break_from_element(element):
    """检查element中是否包含分页符或分节符"""
    try:
        br_elements = element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br')
        for br in br_elements:
            br_type = br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
            if br_type in ['page', 'nextPage', 'evenPage', 'oddPage']:
                return True

        sect_pr = element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
        if sect_pr is not None:
            return True

        return False
    except Exception:
        return False


def _extract_font_info_from_element(element, style_cache):
    """直接从element中提取字体信息"""
    try:
        first_run = element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
        if first_run is None:
            return None, None

        rpr = first_run.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
        if rpr is None:
            return None, None

        font_size_elem = rpr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz')
        font_size = None
        if font_size_elem is not None:
            sz_val = font_size_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            if sz_val:
                font_size = int(sz_val) / 2

        font_name_elem = rpr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
        font_name = None
        if font_name_elem is not None:
            font_name = font_name_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii')
            if not font_name:
                font_name = font_name_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi')

        return font_size, font_name
    except Exception:
        return None, None


def _get_most_common(items: List[Any]) -> Any:
    """获取列表中出现次数最多的元素"""
    if not items:
        return None
    counter = Counter(items)
    return counter.most_common(1)[0][0]


def _build_style_cache(document) -> Dict[str, Dict]:
    """预构建样式缓存字典，提高性能"""
    style_cache = {
        'paragraph_styles': {},
        'run_styles': {},
        'default_fonts': {}
    }

    try:
        for style in document.styles:
            if style.type == 1:  # 段落样式
                style_name = style.name.lower()
                style_info = {
                    'name': style.name,
                    'is_heading': any(heading in style_name for heading in ['heading', 'title', 'header']),
                    'is_centered': False,
                    'font_size': None,
                    'font_name': None,
                    'bold': False
                }

                if hasattr(style, 'paragraph_format') and style.paragraph_format.alignment:
                    style_info['is_centered'] = style.paragraph_format.alignment == 1

                if hasattr(style, 'font'):
                    if style.font.size:
                        style_info['font_size'] = style.font.size.pt
                    if style.font.name:
                        style_info['font_name'] = style.font.name
                    if style.font.bold is not None:
                        style_info['bold'] = style.font.bold

                style_cache['paragraph_styles'][style.name] = style_info

            elif style.type == 2:  # 字符样式
                style_name = style.name.lower()
                style_info = {
                    'name': style.name,
                    'is_bold': any(bold in style_name for bold in ['heading', 'title', 'header', 'strong', 'bold']),
                    'font_size': None,
                    'font_name': None
                }

                if hasattr(style, 'font'):
                    if style.font.size:
                        style_info['font_size'] = style.font.size.pt
                    if style.font.name:
                        style_info['font_name'] = style.font.name
                    if style.font.bold is not None:
                        style_info['is_bold'] = style.font.bold

                style_cache['run_styles'][style.name] = style_info

        if hasattr(document.styles, 'default'):
            default_style = document.styles.default
            if hasattr(default_style, 'font'):
                style_cache['default_fonts'] = {
                    'font_size': default_style.font.size.pt if default_style.font.size else None,
                    'font_name': default_style.font.name if default_style.font.name else None
                }

    except Exception as e:
        print(f"构建样式缓存时出错: {e}")

    return style_cache


def _normalize_style_name(name):
    """标准化样式名称，提高匹配成功率"""
    if not name:
        return ""
    normalized = name.lower()
    normalized = normalized.replace('_', ' ').replace('-', ' ').replace('.', ' ')
    normalized = ' '.join(normalized.split())
    return normalized


def _get_style_name_by_id(document, style_id):
    """通过样式ID获取样式名称"""
    if not style_id:
        return None

    for style in document.styles:
        if style.type == 1:  # 段落样式
            if hasattr(style, '_element') and style._element is not None:
                style_elem_id = style._element.get(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId')
                if style_elem_id == style_id:
                    return style.name

    return None


def _fuzzy_match_style(style_name, cached_styles):
    """模糊匹配样式，支持多种匹配策略"""
    if not style_name:
        return None

    if style_name in cached_styles:
        return cached_styles[style_name]

    normalized_name = style_name.replace(' ', '')
    for cached_name, style_info in cached_styles.items():
        if cached_name.replace(' ', '') == normalized_name:
            return style_info

    normalized_style_name = _normalize_style_name(style_name)
    for cached_name, style_info in cached_styles.items():
        if _normalize_style_name(cached_name) == normalized_style_name:
            return style_info

    return None


def _is_mostly_bold_from_element(element, style_cache, document=None):
    """检查element中大部分run是否加粗（包括样式继承）"""
    try:
        runs = element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
        if not runs:
            return False

        paragraph_bold = False
        ppr = element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if ppr is not None:
            pstyle_elem = ppr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
            if pstyle_elem is not None:
                style_id = pstyle_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                if style_id and 'paragraph_styles' in style_cache:
                    style_name = None
                    if document:
                        style_name = _get_style_name_by_id(document, style_id)

                    if style_name:
                        style_info = _fuzzy_match_style(style_name, style_cache['paragraph_styles'])
                        if style_info and style_info.get('bold', False):
                            paragraph_bold = True

        if paragraph_bold:
            return True

        bold_count = 0
        for run in runs:
            rpr = run.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
            is_bold = False

            if rpr is not None:
                b_elem = rpr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}b')
                if b_elem is not None:
                    is_bold = True

            if is_bold:
                bold_count += 1

        return bold_count >= len(runs) / 2
    except Exception:
        return False


def _is_centered_from_element(element, style_cache, document=None):
    """检查element是否居中（包括样式继承）"""
    try:
        ppr = element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if ppr is not None:
            jc_elem = ppr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}jc')
            if jc_elem is not None:
                jc_val = jc_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                if jc_val == 'center':
                    return True

            pstyle_elem = ppr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
            if pstyle_elem is not None:
                style_id = pstyle_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                if style_id and 'paragraph_styles' in style_cache:
                    style_name = None
                    if document:
                        style_name = _get_style_name_by_id(document, style_id)

                    if style_name:
                        style_info = _fuzzy_match_style(style_name, style_cache['paragraph_styles'])
                        if style_info and style_info.get('is_centered', False):
                            return True

        return False
    except Exception:
        return False


def _has_font_difference_from_element(element, main_font_size, main_font_name, style_cache, document=None):
    """检查element的字体是否与主流字体不同（包括样式继承）"""
    try:
        font_size, font_name = _extract_font_info_from_element(element, style_cache)

        if not font_size or not font_name:
            ppr = element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
            if ppr is not None:
                pstyle_elem = ppr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
                if pstyle_elem is not None:
                    style_id = pstyle_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    if style_id and 'paragraph_styles' in style_cache:
                        style_name = None
                        if document:
                            style_name = _get_style_name_by_id(document, style_id)

                        if style_name:
                            style_info = _fuzzy_match_style(style_name, style_cache['paragraph_styles'])
                            if style_info:
                                if not font_size and style_info.get('font_size'):
                                    font_size = style_info['font_size']
                                if not font_name and style_info.get('font_name'):
                                    font_name = style_info['font_name']

        if font_size and main_font_size and abs(font_size - main_font_size) > 0.5:
            return True

        if font_name and main_font_name and font_name != main_font_name:
            return True

        return False
    except Exception:
        return False


def _generate_markers_from_element(element, main_font_size, main_font_name, style_cache, document=None):
    """从element中生成格式标记（包括样式继承）"""
    markers = []

    try:
        if _is_mostly_bold_from_element(element, style_cache, document):
            markers.append("<加粗>")

        if _is_centered_from_element(element, style_cache, document):
            markers.append("<居中>")

        if _has_font_difference_from_element(element, main_font_size, main_font_name, style_cache, document):
            markers.append("<注意>")

        return "".join(markers)
    except Exception:
        return ""


def _get_overlap_paragraphs(paragraphs: List[Dict], previous_block: List[Dict], win_size: int) -> List[Dict]:
    """获取重叠窗口段落，从前一个块的末尾往前取win_size个字符的段落"""
    if not previous_block:
        return []

    last_para_index = previous_block[-1]['index']
    last_para_position = None

    for i, para in enumerate(paragraphs):
        if para['index'] == last_para_index:
            last_para_position = i
            break

    if last_para_position is None:
        return []

    overlap_paragraphs = []
    char_count = 0

    for i in range(last_para_position, -1, -1):
        para = paragraphs[i]
        if char_count + para['char_count'] <= win_size:
            overlap_paragraphs.insert(0, para)
            char_count += para['char_count']
        else:
            break

    return overlap_paragraphs


def _split_into_blocks(paragraphs: List[Dict], threshold: int, win_size: int) -> List[List[Dict]]:
    """将段落按阈值切块，实现重叠窗口效果"""
    blocks = []
    current_block = []
    current_char_count = 0

    for para in paragraphs:
        char_count = para['char_count']

        if current_char_count + char_count > threshold and current_block:
            blocks.append(current_block)

            overlap_paragraphs = _get_overlap_paragraphs(paragraphs, current_block, win_size)
            current_block = overlap_paragraphs
            current_char_count = sum(p['char_count'] for p in current_block)

            current_block.append(para)
            current_char_count += char_count
        else:
            current_block.append(para)
            current_char_count += char_count

    if current_block:
        blocks.append(current_block)

    return blocks


def process_docx_document(docx_path: str, threshold: int = 1000, win_size: int = 100) -> Dict:
    """将DOCX文档按阈值切块，并为每个段落添加格式标记"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx模块未安装，无法处理DOCX文件")

    print("开始处理DOCX文档...")
    start_time = time.time()

    doc_start = time.time()
    document = delete_auto_numbering_in_docx(docx_path)
    doc_time = time.time() - doc_start
    print(f"文档加载耗时: {doc_time:.3f}秒")

    preprocess_start = time.time()
    document = preprocess_docx_breaks(document)
    preprocess_time = time.time() - preprocess_start
    print(f"预处理耗时: {preprocess_time:.3f}秒")

    cache_start = time.time()
    style_cache = _build_style_cache(document)
    cache_time = time.time() - cache_start
    print(f"样式缓存构建耗时: {cache_time:.3f}秒")

    paragraphs_info = []
    font_sizes = []
    font_names = []

    scan1_start = time.time()
    for element_idx, element in enumerate(document._body._element):
        if element.tag.endswith('p'):
            para_text = _extract_text_from_element(element)

            if not para_text.strip():
                has_page_break = _has_page_break_from_element(element)
                if not has_page_break:
                    continue
                else:
                    print(f"  保留分页符段落 {element_idx}")
                    para_text = "<br>"

            font_size, font_name = _extract_font_info_from_element(element, style_cache)
            if font_size:
                font_sizes.append(font_size)
            if font_name:
                font_names.append(font_name)

            paragraphs_info.append({
                'index': element_idx,
                'text': para_text,
                'element': element,
                'char_count': len(para_text)
            })

        elif element.tag.endswith('tbl'):
            print(f"  跳过表格元素 {element_idx}")
            continue

    scan1_time = time.time() - scan1_start
    print(f"第一遍扫描耗时: {scan1_time:.3f}秒 (处理了 {len(paragraphs_info)} 个段落)")

    font_calc_start = time.time()
    main_font_size = _get_most_common(font_sizes) if font_sizes else None
    main_font_name = _get_most_common(font_names) if font_names else None
    font_calc_time = time.time() - font_calc_start
    print(f"字体统计计算耗时: {font_calc_time:.3f}秒")

    print(f"主流字体大小: {main_font_size}pt")
    print(f"主流字体名称: {main_font_name}")

    processed_paragraphs = []

    scan2_start = time.time()
    for para_info in paragraphs_info:
        element = para_info['element']
        text = para_info['text']
        para_idx = para_info['index']

        markers = _generate_markers_from_element(element, main_font_size, main_font_name, style_cache, document)

        if markers:
            formatted_text = f"{markers} {text}"
        else:
            formatted_text = text

        processed_paragraphs.append({
            'index': para_idx,
            'text': formatted_text,
            'char_count': len(text)
        })

    scan2_time = time.time() - scan2_start
    print(f"第二遍扫描耗时: {scan2_time:.3f}秒 (处理了 {len(processed_paragraphs)} 个段落)")

    block_start = time.time()
    blocks = _split_into_blocks(processed_paragraphs, threshold, win_size)
    print("\n切块结果展示:")
    print("=" * 60)
    for block_idx, block in enumerate(blocks):
        print(f"\n块 {block_idx + 1} (包含 {len(block)} 个段落):")
        print("-" * 40)
        total_chars = sum(para['char_count'] for para in block)
        print(f"总字符数: {total_chars}")

        for para_idx, para_info in enumerate(block):
            para_text = para_info['text']
            display_text = para_text[:80] + "..." if len(para_text) > 80 else para_text
            print(f"  {para_idx + 1}. 段落{para_info['index']}: {display_text}")

        print("-" * 40)
    block_time = time.time() - block_start
    print(f"切块处理耗时: {block_time:.3f}秒 (生成了 {len(blocks)} 个块)")

    format_start = time.time()
    result = []

    instruction_template = "提供一个文档的段落片段，按每个段落的格式设置了四个标签，<加粗>表示大部分段落是加粗的，<居中>表示该段落是居中的，<注意>表示该段落的字体字号与大部分正文段落不同，需要注意。请根据以上信息，尝试找出标题段落的序号及段落内容，参考以下格式输出，不要输出其他思考内容：{\"32\": \"<加粗> 一、投标函\", \"57\": \"<居中> 二、法人身份证\"...}"

    for block_idx, block in enumerate(blocks):
        block_dict = {}
        for para_info in block:
            block_dict[str(para_info['index'])] = para_info['text']

        data_item = f""""instruction": {instruction_template},"input": {str(block_dict)}"""

        result.append(data_item)

    format_time = time.time() - format_start
    print(f"格式转换耗时: {format_time:.3f}秒")
    print(f"生成了 {len(result)} 个训练数据项")

    total_paragraphs = int(paragraphs_info[-1]['index']) + 1
    print(f"文档总段数: {total_paragraphs}")

    # 生成JSON格式的数据并保存
    json_data = []
    for block_idx, block in enumerate(blocks):
        block_dict = {}
        for para_info in block:
            block_dict[str(para_info['index'])] = para_info['text']
        
        json_item = {
            "instruction": instruction_template,
            "input": str(block_dict),
            "output": "{}"
        }
        json_data.append(json_item)
    
    try:
        input_dir = os.path.dirname(docx_path)
        json_filename = "block_with_title.json"
        json_path = os.path.join(input_dir, json_filename)
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)
        print(f"已保存JSON数据到: {json_path}")
    except Exception as e:
        print(f"保存JSON数据时出错: {e}")

    total_time = time.time() - start_time
    print(f"\n总处理时间: {total_time:.3f}秒")

    return {
        "result": result,
        "total_paragraphs": total_paragraphs,
        "document_info": {
            "docx_path": docx_path,
            "threshold": threshold,
            "win_size": win_size,
            "total_blocks": len(result)
        }
    }


# =============================================================================
# 大模型调用相关函数（需要配置）
# =============================================================================

# 注意：这些函数需要配置API密钥和URL，这里提供占位符实现
# 实际使用时需要根据实际情况配置


def lora_general_caller(query: str, switch: int, model_name: str=LORA_MODEL_NAME, stream: bool=False, top_p: float=0.3, temperature: float=0.2, extra_body: dict=None,
                        retries=2):
    try:
        start_time = time.time()
        reasoning_content = ""  # 定义完整思考过程
        answer_content = ""  # 定义完整回复
        is_answering = False  # 判断是否结束思考过程并开始回复
        if switch == 1:
            model_name = LORA_MODEL_NAME

        messages = [{"role": "user", "content": query}]

        # 创建聊天完成请求
        completion = call_llm(messages, switch, model_name=model_name, stream=stream, top_p=top_p,
                              temperature=temperature, extra_body=extra_body)

        # =============== 新增非流式处理逻辑 ===============
        if not stream:

            # 处理无有效choices的情况
            if not completion.choices:
                if hasattr(completion, 'usage'):
                    print("\nUsage:")
                    print(completion.usage)
                return ""

            # 获取完整的assistant消息
            assistant_message = completion.choices[0].message

            # 提取思考过程（如果存在）
            if hasattr(assistant_message, 'reasoning_content') and assistant_message.reasoning_content:
                reasoning_content = assistant_message.reasoning_content

            # 提取回复内容（如果存在）
            if hasattr(assistant_message, 'content') and assistant_message.content:
                answer_content = assistant_message.content

        # =============== 流式处理逻辑（保持不变） ===============
        else:
            for chunk in completion:
                # 如果chunk.choices为空，则打印usage
                if not chunk.choices:
                    print("\nUsage:")
                    print(chunk.usage)
                else:
                    delta = chunk.choices[0].delta
                    # 处理思考过程
                    if hasattr(delta, 'reasoning_content') and delta.reasoning_content:
                        reasoning_content += delta.reasoning_content
                    else:
                        # 开始回复
                        if delta.content and is_answering is False:
                            is_answering = True
                        # 收集回复内容
                        if delta.content:
                            answer_content += delta.content

        # 统一的后处理逻辑
        if "</think>" in answer_content:
            answer_content = answer_content.split("</think>")[1].strip()

        cost_time = time.time() - start_time
        return answer_content

    except (RateLimitError, APIConnectionError, APIError) as e:
        if retries > 0:
            wait_time = (3 - retries)
            print(f"Request failed, retrying in {wait_time} seconds...")
            time.sleep(wait_time)
            return lora_general_caller(query, switch, model_name, stream, top_p, temperature, extra_body, retries - 1)
        return ""
    except Exception as e:
        print(f"An error occurred: {e}")
        return ""


def call_llm(messages, switch, model_name=LORA_MODEL_NAME, stream=False, top_p=0.8, temperature=1.0, extra_body=None, presence_penalty=None):
    print(f"当前调用模型：{model_name}")

    kwargs = {
        "model": model_name,
        "messages": messages,
        "stream": stream
    }
    if extra_body:
        kwargs["extra_body"] = extra_body
    if presence_penalty:
        kwargs["presence_penalty"] = presence_penalty

    return lora_client.chat.completions.create(**kwargs)


def lora_call_request(query, switch=0):
    """通用lora_request，代码复用"""
    with ThreadPoolExecutor(max_workers=1) as executor:
        future = executor.submit(lora_general_caller, query, switch)
        try:
            ans = future.result(timeout=200)
        except TimeoutError:
            ans = ""
    return ans


def call_request(query, model="", extra_body_mode=0):
    """
    通用request，代码复用（占位符实现）
    实际使用时需要配置API密钥和URL
    """
    if not OPENAI_AVAILABLE:
        print("警告: openai模块未安装，无法调用大模型")
        return ""
    
    # 这里需要根据实际情况配置
    print(f"警告: call_request未配置，返回空字符串")
    print(f"查询内容: {query[:100]}...")
    return ""


def ask_for_title_level(title_content):
    """获取标题层级"""
    query = f"""
**任务描述：**
分析提供的文档标题文本，识别每个标题的层级结构及其从属关系。输出一个 JSON 对象，其中键是标题的原始序号（在输入文本中的顺序号），值是其对应的层级编号。

**输入文本：**
\"\"\"{str(title_content)}\"\"\"

**处理规则：**
1.  **起始标题处理：**
    *   如果输入文本中的**第一个**标题（序号 1）符合类似 `第X章 XXX`、`第X节 XXX` 或 `第X条 XXX` 的格式（其中 X 是数字或汉字数字），则**跳过**该标题，**不**将其包含在输出 JSON 中。处理从第二个标题（序号 2）开始。
    *   如果第一个标题不符合上述格式，则正常处理。

2.  **层级判定原则：**
    *   主要依据标题前的**前缀编号样式**（如 `1.`, `1.1`, `(1)`, `一、`, `（一）` 等）来判断层级。
    *   **一般规则：** 相同或相似前缀样式的标题通常属于同一层级。例如，所有以 `1.`, `2.`, `3.` 开头的标题属于同一层级（一级）；所有以 `1.1`, `1.2`, `2.1` 开头的标题属于下一层级（二级）。
    *   **实际为准：** 如果遇到前缀样式变化但逻辑上应属于同一层级，或样式相同但逻辑上应分属不同层级的情况，**以标题内容在文档中实际表达的层级关系为准**。仔细分析编号的连续性和缩进（如果输入包含缩进信息）。

3.  **无前缀标题处理：**
    *   **开头无前缀标题：** 如果在处理开始后（见规则1），文档**开头部分**连续出现一个或多个**没有任何前缀编号样式**的标题：
        *   这些标题被视为**"编外标题"**。
        *   它们被赋予与**一级标题**相同的层级（即层级编号从 `1` 开始计数）。
        *   后续出现的第一个带编号样式的标题，如果其样式表明它是一级标题，则与这些"编外标题"**属于同一层级**（即也是层级 `1`, `2`, `3`...）。
    *   **中间无前缀标题：** 如果文档**中间**出现一个**没有任何前缀编号样式**的标题：
        *   它通常被视为**上一标题的直属下级标题**（即层级编号在上一标题层级编号后增加一级，如上一标题是 `1`，则它为 `1-1`；上一标题是 `2-3`，则它为 `2-3-1`）。
        *   如果上下文强烈暗示它应属于更高层级，则以实际内容为准（此情况较少见）。

4.  **层级编号规则：**
    *   层级编号采用 **`父级编号` + `-` + `子级序号`** 的格式。
    *   **顶级标题（一级）：** 使用整数编号 `1`, `2`, `3`, ... 表示。
    *   **子级标题：**
        *   一级标题下的第一子级（二级）：`1-1`, `1-2`, `1-3`, ... `2-1`, `2-2`, ...
        *   二级标题下的第一子级（三级）：`1-1-1`, `1-1-2`, ... `2-3-1`, ...
        *   依此类推。
    *   **编号起点：** 每个层级的序号都从 `1` 开始重新计数。
    *   **"编外标题"编号：** 被视为一级标题，编号为 `1`, `2`, `3`, ...

**输出要求：**
*   格式：**严格的 JSON 对象**。
*   键 (Key)：标题在输入文本中的**原始序号**（字符串类型，如 `"33"`）。
*   值 (Value)：该标题对应的**层级编号**（字符串类型，如 `"1"`, `"1-1"`, `"2-3-4"`）。
*   **仅包含：** 根据规则处理后的标题条目。被跳过的标题（规则1）不应出现在输出中。
*   **示例输出格式：**
{{
    "33": "1",      // 一级标题
    "34": "1-1",    // 一级标题下的第一个二级标题
    "35": "1-2",    // 一级标题下的第二个二级标题
    "47": "2",      // 下一个一级标题
    "48": "2-1",    // 第二个一级标题下的第一个二级标题
    "50": "3",      // 一个"编外标题"（被视为一级）
    "51": "3-1"     // 上一个"编外标题"下的子标题
}}
"""
    title_level_ans = call_request(query)
    title_level_json = extract_json_from_string(title_level_ans)
    return title_level_json


# =============================================================================
# 文档结构组织相关函数
# =============================================================================

def extract_original_title(formatted_content: str) -> str:
    """从格式化的标题内容中提取原始标题"""
    original = re.sub(r'<[^>]+>', '', formatted_content)
    original = ' '.join(original.split())
    return original.strip()


def calculate_title_level(original_level: str) -> int:
    """计算标题的层级深度（数字表示）"""
    try:
        if not original_level or original_level == "0":
            return 1
        depth = original_level.count('-') + 1
        return max(1, depth)
    except Exception:
        return 1


def is_child_title(parent_level: str, child_level: str) -> bool:
    """判断child_level是否是parent_level的子标题"""
    try:
        if not parent_level or not child_level:
            return False
        if child_level.startswith(parent_level + "-"):
            return True
        return False
    except Exception:
        return False


def calculate_title_end_paragraph(current_para: int, current_level: str, all_para_nums: List[int], 
                                  para_to_title: Dict, total_paragraphs: int) -> int:
    """计算标题的结束段落，实现层级包含关系"""
    try:
        following_titles = []
        for para_num in all_para_nums:
            if para_num > current_para:
                title_info = para_to_title[para_num]
                title_level = title_info.get("level", "0")
                following_titles.append((para_num, title_level))

        if not following_titles:
            return total_paragraphs - 1

        current_depth = calculate_title_level(current_level)

        for para_num, title_level in following_titles:
            title_depth = calculate_title_level(title_level)
            if title_depth <= current_depth:
                return para_num - 1

        return total_paragraphs - 1
    except Exception as e:
        print(f"计算标题结束段落时出错: {e}")
        return current_para


def is_title_leaf_block(current_para: int, current_level: str, all_para_nums: List[int], 
                        para_to_title: Dict) -> bool:
    """判断标题是否是叶子块（没有子标题）"""
    try:
        following_titles = []
        for para_num in all_para_nums:
            if para_num > current_para:
                title_info = para_to_title[para_num]
                title_level = title_info.get("level", "0")
                following_titles.append((para_num, title_level))

        if not following_titles:
            return True

        current_depth = calculate_title_level(current_level)

        for para_num, title_level in following_titles:
            title_depth = calculate_title_level(title_level)
            if title_depth <= current_depth:
                return True
            if title_depth > current_depth:
                if is_child_title(current_level, title_level):
                    return False

        return True
    except Exception as e:
        print(f"判断叶子块时出错: {e}")
        return True


def check_block_has_table(start_para: int, end_para: int, document) -> bool:
    """检查块范围内是否包含表格"""
    try:
        if not document or not DOCX_AVAILABLE:
            return False

        for i, element in enumerate(document._body._element):
            if element.tag.endswith('p'):
                if start_para <= i <= end_para:
                    next_element = element.getnext()
                    if next_element is not None and next_element.tag.endswith('tbl'):
                        return True
            elif element.tag.endswith('tbl'):
                if start_para <= i <= end_para:
                    return True

        return False
    except Exception as e:
        print(f"检查表格时出错: {e}")
        return False


def build_parallel_blocks(sorted_titles: List[Tuple[str, Dict]], total_paragraphs: int, 
                          result: List[str], document) -> List[Dict]:
    """构建并列的块结构，每个标题都是独立的dict"""
    try:
        blocks = []

        para_to_title = {}
        for para_num, title_info in sorted_titles:
            para_to_title[int(para_num)] = title_info

        all_para_nums = sorted(para_to_title.keys())

        for i, (para_num, title_info) in enumerate(sorted_titles):
            try:
                para_num_int = int(para_num)
                original_level = title_info.get("level", "0")
                content = title_info.get("content", "")

                original_title = extract_original_title(content)
                level = calculate_title_level(original_level)

                start_para = para_num_int
                end_para = calculate_title_end_paragraph(para_num_int, original_level, all_para_nums, 
                                                         para_to_title, total_paragraphs)

                is_leaf_block = is_title_leaf_block(para_num_int, original_level, all_para_nums, para_to_title)

                has_table = check_block_has_table(start_para, end_para, document)

                block = {
                    "title_idx": original_level,
                    "title_content": original_title,
                    "start_para": start_para,
                    "end_para": end_para,
                    "level": level,
                    "is_leaf_block": is_leaf_block,
                    "has_table": has_table,
                    "content_summary": "",
                    "insert_resource": []
                }

                blocks.append(block)

            except Exception as e:
                print(f"处理段落 {para_num} 时出错: {e}")
                continue

        return blocks

    except Exception as e:
        print(f"构建并列块结构时出错: {e}")
        traceback.print_exc()
        return []


def organize_document_structure(final_dict: Dict, process_result: Dict, document) -> Dict:
    """将final_dict重新组织为更清晰的文档分块结构"""
    try:
        result = process_result.get("result", [])
        total_paragraphs = process_result.get("total_paragraphs", 0)

        print(f"文档总段数: {total_paragraphs}")

        valid_titles = {}
        for k, v in final_dict.items():
            try:
                if isinstance(v, dict) and v.get("level") != "-1":
                    valid_titles[k] = v
            except Exception as e:
                print(f"处理标题 {k} 时出错: {e}")
                continue

        if not valid_titles:
            print("警告: 没有找到有效的标题")
            return {"blocks": []}

        try:
            sorted_titles = sorted(valid_titles.items(), key=lambda x: int(x[0]))
        except ValueError as e:
            print(f"排序段落号时出错: {e}")
            sorted_titles = list(valid_titles.items())

        blocks = build_parallel_blocks(sorted_titles, total_paragraphs, result, document)

        return {"blocks": blocks}

    except Exception as e:
        print(f"组织文档结构时出错: {e}")
        traceback.print_exc()
        return {"blocks": []}


# =============================================================================
# 文件转换相关函数（需要配置）
# =============================================================================

def pdf_to_docx(pdf_path: str) -> str:
    """
    将PDF转换为DOCX（占位符实现）
    实际使用时需要配置转换服务的URL
    """
    if not REQUESTS_AVAILABLE:
        print("警告: requests模块未安装，无法进行文件转换")
        return ""
    
    # 这里需要根据实际情况配置转换服务
    # 示例：
    # url = "your_pdf_to_docx_url"
    # with open(pdf_path, "rb") as f:
    #     files = {"file": (os.path.basename(pdf_path), f.read())}
    #     response = requests.post(url, files=files, timeout=None)
    #     if response.status_code == 200:
    #         # 处理响应，保存文件
    #         return converted_file_path
    
    print(f"警告: pdf_to_docx未配置，无法转换文件: {pdf_path}")
    return ""


def doc_to_docx(doc_path: str) -> str:
    """
    将DOC转换为DOCX（占位符实现）
    实际使用时需要配置转换服务的URL
    """
    if not REQUESTS_AVAILABLE:
        print("警告: requests模块未安装，无法进行文件转换")
        return ""
    
    # 这里需要根据实际情况配置转换服务
    print(f"警告: doc_to_docx未配置，无法转换文件: {doc_path}")
    return ""


def preprocess_input_file(input_file_path, output_dir):
    """
    预处理输入文件：判断文件类型并统一转换为docx格式
    """
    print(f"\n预处理: 判断文件类型并转换...")
    print(f"输入文件: {input_file_path}")
    
    if not os.path.exists(input_file_path):
        raise FileNotFoundError(f"输入文件不存在: {input_file_path}")
    
    file_ext = Path(input_file_path).suffix.lower()
    file_type = ZB_FILE_TYPE_CONFIG.get(file_ext, ZBFileType.UNKNOWN)
    
    if file_type == ZBFileType.UNKNOWN:
        raise ValueError(f"不支持的文件类型: {file_ext}。支持的类型: .pdf, .doc, .docx")
    
    if file_type == ZBFileType.DOCX:
        print(f"文件已经是docx格式，无需转换")
        return input_file_path
    
    input_filename = Path(input_file_path).stem
    converted_file_path = os.path.join(output_dir, f"{input_filename}.docx")
    
    if os.path.exists(converted_file_path):
        print(f"转换后的文件已存在: {converted_file_path}")
        return converted_file_path
    
    try:
        if file_type == ZBFileType.PDF:
            print(f"正在将PDF转换为DOCX...")
            converted_path = pdf_to_docx(input_file_path)
            if not converted_path:
                raise Exception("PDF转DOCX失败，返回路径为空")
            
        elif file_type == ZBFileType.DOC:
            print(f"正在将DOC转换为DOCX...")
            converted_path = doc_to_docx(input_file_path)
            if not converted_path:
                raise Exception("DOC转DOCX失败，返回路径为空")
        
        if os.path.dirname(os.path.abspath(converted_path)) != os.path.abspath(output_dir):
            if os.path.exists(converted_path):
                shutil.move(converted_path, converted_file_path)
                converted_path = converted_file_path
            else:
                raise FileNotFoundError(f"转换后的文件不存在: {converted_path}")
        
        print(f"文件转换完成: {converted_path}")
        return converted_path
        
    except Exception as e:
        print(f"文件转换失败: {str(e)}")
        traceback.print_exc()
        raise


# =============================================================================
# 主处理函数
# =============================================================================

def extract_titles_with_lora(intermediate_docx_path, threshold=1000, win_size=200):
    """对中间文件应用lora模型找标题"""
    print(f"\n开始对中间文件应用lora模型找标题: {intermediate_docx_path}")
    
    print("加载并预处理文档...")
    document = delete_auto_numbering_in_docx(intermediate_docx_path)
    document = preprocess_docx_breaks(document)
    
    print("处理文档，提取段落信息...")
    process_result = process_docx_document(intermediate_docx_path, threshold, win_size)
    
    result = process_result.get("result", [])
    total_paragraphs = process_result.get("total_paragraphs", 0)
    
    print(f"\n文档处理完成！")
    print(f"处理结果统计:")
    print(f"  - 总数据项数: {len(result)}")
    print(f"  - 文档总段数: {total_paragraphs}")
    
    final_dict = {}
    lora_structure = {}
    ans_check = {}
    
    print(f"\n开始并行处理 {len(result)} 个数据项...")
    
    def process_data_item(index_data):
        """处理单个数据项的辅助函数"""
        i, data_item = index_data
        try:
            print(f"处理数据项 {i + 1}/{len(result)}...")
            lora_ans = lora_call_request(data_item, switch=1)
            print(f"  数据项 {i + 1} Lora输出完成")
            lora_ans_json = extract_json_from_string(lora_ans)
            return i, lora_ans_json
        except Exception as e:
            print(f"  数据项 {i + 1} 处理失败: {str(e)}")
            return i, {}
    
    with ThreadPoolExecutor(max_workers=20) as executor:
        futures = {executor.submit(process_data_item, (i, data_item)): i
                   for i, data_item in enumerate(result)}
        
        results_dict = {}
        for future in as_completed(futures):
            try:
                i, lora_ans_json = future.result()
                results_dict[i] = lora_ans_json
            except Exception as e:
                original_index = futures[future]
                print(f"任务 {original_index} 执行出错: {str(e)}")
                results_dict[original_index] = {}
    
    print("\n整合处理结果...")
    for i in range(len(result)):
        if i in results_dict:
            lora_ans_json = results_dict[i]
            ans_check[i] = lora_ans_json
            if isinstance(lora_ans_json, dict):
                for k, v in lora_ans_json.items():
                    lora_structure[str(k)] = str(v)
    
    # 更新block_with_title.json中的output字段
    try:
        json_dir = os.path.dirname(intermediate_docx_path)
        json_path = os.path.join(json_dir, "block_with_title.json")
        
        if os.path.exists(json_path):
            print(f"\n更新block_with_title.json中的output字段...")
            with open(json_path, "r", encoding="utf-8") as f:
                json_data = json.load(f)
            
            for i in range(len(json_data)):
                if i in results_dict:
                    lora_ans_json = results_dict[i]
                    if isinstance(lora_ans_json, dict):
                        output_str = str(lora_ans_json)
                        json_data[i]["output"] = output_str
                    else:
                        json_data[i]["output"] = "{}"
                else:
                    json_data[i]["output"] = "{}"
            
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(json_data, f, ensure_ascii=False, indent=2)
            print(f"已更新block_with_title.json，共更新 {len(json_data)} 个块的output字段")
        else:
            print(f"警告: 未找到block_with_title.json文件: {json_path}")
    except Exception as e:
        print(f"更新block_with_title.json时出错: {e}")
        traceback.print_exc()
    
    # 获取标题层级
    print("\n获取标题层级...")
    title_structure = ask_for_title_level(lora_structure)
    
    # 构建final_dict
    for k, v in lora_structure.items():
        final_dict[k] = {}
        final_dict[k]["content"] = v
        if k in title_structure.keys():
            final_dict[k]["level"] = title_structure[k]
        else:
            final_dict[k]["level"] = "-1"
    
    print("\n构建的final_dict:")
    print(json.dumps(final_dict, ensure_ascii=False, indent=2))
    
    # 重新组织数据结构为更清晰的分块结构
    print("\n重新组织文档结构...")
    structured_data = organize_document_structure(final_dict, process_result, document)
    print("\n重新组织后的文档结构:")
    print(json.dumps(structured_data, ensure_ascii=False, indent=2))
    
    return {
        "lora_structure": lora_structure,
        "final_dict": final_dict,
        "structured_data": structured_data,
        "title_structure": title_structure,
        "results_dict": results_dict
    }


def main(input_file_path, output_dir, threshold=1000, win_size=200):
    """
    主函数
    
    Args:
        input_file_path: 招标文件路径（支持pdf、doc、docx格式）
        output_dir: 输出目录，所有结果将保存在此目录中
        threshold: 字符数阈值，超过此值则切块（默认1000）
        win_size: 窗口大小，在块开头往前取的字符数（默认200）
    """
    output_dir = os.path.abspath(output_dir)
    os.makedirs(output_dir, exist_ok=True)
    
    print(f"输出目录: {output_dir}")
    
    try:
        # 步骤2: 对中间文件应用lora模型找标题
        print(f"\n步骤2: 对中间文件应用lora模型找标题...")
        title_results = extract_titles_with_lora(
            input_file_path,
            threshold=threshold,
            win_size=win_size
        )
        
        # 步骤3: 保存最终的标题json
        print(f"\n步骤3: 保存最终的标题json...")
        
        output_filename = os.path.splitext(os.path.basename(input_file_path))[0] + "_title_result.json"
        output_path = os.path.join(output_dir, output_filename)
        
        output_data = title_results["lora_structure"]
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, ensure_ascii=False, indent=2)
        
        print(f"最终标题json已保存到: {output_path}")
        
        print("\n处理完成！")
        
    except Exception as e:
        print(f"处理过程中出现错误: {str(e)}")
        traceback.print_exc()
        return


if __name__ == "__main__":
    # 请在此处修改参数进行调试
    main(
        input_file_path=r"data/3.docx",  # 请填入招标文件路径（支持pdf、doc、docx）
        output_dir=r"data",  # 请填入输出目录
        threshold=1000,  # 可选，默认1000
        win_size=200  # 可选，默认200
    )

