"""
图片检测模块
提供图片检测、印章识别等功能
"""

import io
from docx import Document
from docx.oxml.ns import qn

# 检查PIL是否可用
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False


def has_image(element):
    """
    检查元素中是否包含图片（优化版）
    
    Args:
        element: XML元素
        
    Returns:
        如果包含图片返回True，否则返回False
    """
    # 使用命名空间常量，避免重复字符串拼接
    ns_drawing = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
    ns_main = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
    
    # 先检查inline元素（更常见）
    if element.find(f'.//{ns_drawing}inline') is not None:
        return True
    
    # 再检查blip元素
    if element.find(f'.//{ns_main}blip') is not None:
        return True
    
    return False


def has_image_in_run(run):
    """
    检查run中是否包含图片
    
    Args:
        run: Run对象
        
    Returns:
        如果包含图片返回True，否则返回False
    """
    ns_drawing = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
    inline = run._element.find(f'.//{ns_drawing}inline')
    return inline is not None


def get_image_data_from_run(doc: Document, run):
    """
    从run中提取图片数据
    
    Args:
        doc: Document对象
        run: Run对象
        
    Returns:
        (图片字节数据, 文件扩展名) 元组，如果无法获取则返回None
    """
    # 检查run中是否有图片
    ns_drawing = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
    inline = run._element.find(f'.//{ns_drawing}inline')
    if inline is not None:
        # 查找图片的relationship ID
        ns_main = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
        blip = inline.find(f'.//{ns_main}blip')
        if blip is not None:
            # 尝试获取嵌入的图片 (r:embed)
            r_embed = blip.get(qn('r:embed'))
            if r_embed:
                try:
                    # 获取图片数据
                    image_part = doc.part.related_parts[r_embed]
                    image_bytes = image_part.blob
                    
                    # 获取图片扩展名
                    content_type = image_part.content_type
                    ext_map = {
                        'image/png': '.png',
                        'image/jpeg': '.jpg',
                        'image/jpg': '.jpg',
                        'image/gif': '.gif',
                        'image/bmp': '.bmp',
                        'image/x-emf': '.emf',
                        'image/x-wmf': '.wmf',
                    }
                    ext = ext_map.get(content_type, '.png')
                    return (image_bytes, ext)
                except (KeyError, AttributeError, Exception):
                    return None
            
            # 链接图片无法获取数据
            r_link = blip.get(qn('r:link'))
            if r_link:
                return None
    
    return None


def get_image_size_from_run(run):
    """
    从run中获取图片尺寸（EMU单位）
    
    Args:
        run: Run对象
        
    Returns:
        (width, height) 元组，单位为EMU，如果无法获取则返回None
    """
    # 检查run中是否有图片
    ns_drawing = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
    inline = run._element.find(f'.//{ns_drawing}inline')
    if inline is not None:
        # 尝试从run的XML中获取图片尺寸
        extent = inline.find(f'.//{ns_drawing}extent')
        if extent is not None:
            cx = extent.get('cx')  # 宽度
            cy = extent.get('cy')  # 高度
            if cx and cy:
                try:
                    return (int(cx), int(cy))
                except ValueError:
                    pass
    
    return None


def is_square_image_in_run(run, aspect_ratio_tolerance=0.03):
    """
    检查run中的图片是否近似方形（长宽比接近1:1）
    
    Args:
        run: Run对象
        aspect_ratio_tolerance: 长宽比容差，默认0.03（即长宽比在0.97-1.03之间认为是方形，3%容差）
        
    Returns:
        如果图片近似方形返回True，否则返回False
    """
    size = get_image_size_from_run(run)
    if size is None:
        return False
    
    width_emu, height_emu = size
    if width_emu == 0 or height_emu == 0:
        return False
    
    # 计算长宽比（宽/高）
    aspect_ratio = width_emu / height_emu
    
    # 如果长宽比接近1:1（在容差范围内），认为是方形
    is_square = abs(aspect_ratio - 1.0) <= aspect_ratio_tolerance
    
    return is_square


def is_red_image_in_run(doc: Document, run, red_threshold=0.3):
    """
    检查run中的图片是否主要是红色
    
    Args:
        doc: Document对象
        run: Run对象
        red_threshold: 红色像素占比阈值，默认0.3（30%以上红色像素认为是红色图片）
        
    Returns:
        如果图片主要是红色返回True，否则返回False
    """
    if not PIL_AVAILABLE:
        return False
    
    # 获取图片数据
    image_data = get_image_data_from_run(doc, run)
    if image_data is None:
        return False
    
    image_bytes, ext = image_data
    
    try:
        # 打开图片
        img = Image.open(io.BytesIO(image_bytes))
        
        # 转换为RGB模式（如果不是）
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        # 获取图片尺寸
        width, height = img.size
        total_pixels = width * height
        
        if total_pixels == 0:
            return False
        
        # 统计红色像素数量
        red_pixel_count = 0
        
        # 采样检测（如果图片太大，采样检测以提高性能）
        sample_step = max(1, min(width, height) // 100)  # 采样步长
        
        for y in range(0, height, sample_step):
            for x in range(0, width, sample_step):
                r, g, b = img.getpixel((x, y))
                
                # 判断是否为红色像素
                # 红色特征：R值较高，且R > G和R > B
                if r > 100 and r > g * 1.2 and r > b * 1.2:
                    red_pixel_count += 1
        
        # 计算红色像素占比
        sampled_pixels = ((width // sample_step) * (height // sample_step))
        red_ratio = red_pixel_count / sampled_pixels if sampled_pixels > 0 else 0
        
        is_red = red_ratio >= red_threshold
        
        return is_red
        
    except Exception:
        return False


def has_stamp_features_in_run(doc: Document, run, check_square=True, check_red=True, 
                               aspect_ratio_tolerance=0.03, red_threshold=0.3):
    """
    检查run中的图片是否具有印章特征（近似方形或红色）
    
    Args:
        doc: Document对象
        run: Run对象
        check_square: 是否检查是否为方形，默认True
        check_red: 是否检查是否为红色，默认True
        aspect_ratio_tolerance: 长宽比容差，默认0.03（3%容差）
        red_threshold: 红色像素占比阈值，默认0.3
        
    Returns:
        (是否具有印章特征, 特征详情字典) 元组
    """
    features = {
        'is_square': False,
        'is_red': False
    }
    
    if check_square:
        features['is_square'] = is_square_image_in_run(run, aspect_ratio_tolerance)
    
    if check_red:
        features['is_red'] = is_red_image_in_run(doc, run, red_threshold)
    
    # 只要满足任一特征就认为是印章
    has_features = features['is_square'] or features['is_red']
    
    return (has_features, features)


def get_image_data_from_paragraph(doc: Document, element, element_to_para=None):
    """
    从段落中提取图片数据
    
    Args:
        doc: Document对象
        element: 段落XML元素
        element_to_para: 元素到Paragraph对象的映射字典（可选，用于性能优化）
        
    Returns:
        (图片字节数据, 文件扩展名) 元组，如果无法获取则返回None
    """
    # 找到对应的Paragraph对象
    if element_to_para is not None:
        para_obj = element_to_para.get(element)
    else:
        para_obj = None
        for para in doc.paragraphs:
            if para._element == element:
                para_obj = para
                break
    
    if para_obj is None:
        return None
    
    # 遍历段落中的所有runs，查找包含图片的run
    for run in para_obj.runs:
        # 检查run中是否有图片
        ns_drawing = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
        inline = run._element.find(f'.//{ns_drawing}inline')
        if inline is not None:
            # 查找图片的relationship ID
            ns_main = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
            blip = inline.find(f'.//{ns_main}blip')
            if blip is not None:
                # 尝试获取嵌入的图片 (r:embed)
                r_embed = blip.get(qn('r:embed'))
                if r_embed:
                    try:
                        # 获取图片数据
                        image_part = doc.part.related_parts[r_embed]
                        image_bytes = image_part.blob
                        
                        # 获取图片扩展名
                        content_type = image_part.content_type
                        ext_map = {
                            'image/png': '.png',
                            'image/jpeg': '.jpg',
                            'image/jpg': '.jpg',
                            'image/gif': '.gif',
                            'image/bmp': '.bmp',
                            'image/x-emf': '.emf',
                            'image/x-wmf': '.wmf',
                        }
                        ext = ext_map.get(content_type, '.png')
                        return (image_bytes, ext)
                    except (KeyError, AttributeError, Exception):
                        return None
                
                # 链接图片无法获取数据
                r_link = blip.get(qn('r:link'))
                if r_link:
                    return None
    
    return None


def get_image_size_from_paragraph(doc: Document, element, element_to_para=None):
    """
    从段落中获取图片尺寸（EMU单位）
    
    Args:
        doc: Document对象
        element: 段落XML元素
        element_to_para: 元素到Paragraph对象的映射字典（可选，用于性能优化）
        
    Returns:
        (width, height) 元组，单位为EMU，如果无法获取则返回None
    """
    # 找到对应的Paragraph对象
    if element_to_para is not None:
        para_obj = element_to_para.get(element)
    else:
        para_obj = None
        for para in doc.paragraphs:
            if para._element == element:
                para_obj = para
                break
    
    if para_obj is None:
        return None
    
    # 遍历段落中的所有runs，查找包含图片的run
    for run in para_obj.runs:
        # 检查run中是否有图片
        ns_drawing = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
        inline = run._element.find(f'.//{ns_drawing}inline')
        if inline is not None:
            # 尝试从run的XML中获取图片尺寸
            extent = inline.find(f'.//{ns_drawing}extent')
            if extent is not None:
                cx = extent.get('cx')  # 宽度
                cy = extent.get('cy')  # 高度
                if cx and cy:
                    try:
                        return (int(cx), int(cy))
                    except ValueError:
                        pass
    
    return None


def is_square_image(doc: Document, element, element_to_para=None, aspect_ratio_tolerance=0.03):
    """
    检查图片是否近似方形（长宽比接近1:1）
    
    Args:
        doc: Document对象
        element: 段落XML元素
        element_to_para: 元素到Paragraph对象的映射字典（可选）
        aspect_ratio_tolerance: 长宽比容差，默认0.03（即长宽比在0.97-1.03之间认为是方形，3%容差）
        
    Returns:
        如果图片近似方形返回True，否则返回False
    """
    size = get_image_size_from_paragraph(doc, element, element_to_para)
    if size is None:
        return False
    
    width_emu, height_emu = size
    if width_emu == 0 or height_emu == 0:
        return False
    
    # 计算长宽比（宽/高）
    aspect_ratio = width_emu / height_emu
    
    # 如果长宽比接近1:1（在容差范围内），认为是方形
    is_square = abs(aspect_ratio - 1.0) <= aspect_ratio_tolerance
    
    return is_square


def is_red_image(doc: Document, element, element_to_para=None, red_threshold=0.3):
    """
    检查图片是否主要是红色
    
    Args:
        doc: Document对象
        element: 段落XML元素
        element_to_para: 元素到Paragraph对象的映射字典（可选）
        red_threshold: 红色像素占比阈值，默认0.3（30%以上红色像素认为是红色图片）
        
    Returns:
        如果图片主要是红色返回True，否则返回False
    """
    if not PIL_AVAILABLE:
        return False
    
    # 获取图片数据
    image_data = get_image_data_from_paragraph(doc, element, element_to_para)
    if image_data is None:
        return False
    
    image_bytes, ext = image_data
    
    try:
        # 打开图片
        img = Image.open(io.BytesIO(image_bytes))
        
        # 转换为RGB模式（如果不是）
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        # 获取图片尺寸
        width, height = img.size
        total_pixels = width * height
        
        if total_pixels == 0:
            return False
        
        # 统计红色像素数量
        red_pixel_count = 0
        
        # 采样检测（如果图片太大，采样检测以提高性能）
        sample_step = max(1, min(width, height) // 100)  # 采样步长
        
        for y in range(0, height, sample_step):
            for x in range(0, width, sample_step):
                r, g, b = img.getpixel((x, y))
                
                # 判断是否为红色像素
                # 红色特征：R值较高，且R > G和R > B
                if r > 100 and r > g * 1.2 and r > b * 1.2:
                    red_pixel_count += 1
        
        # 计算红色像素占比
        sampled_pixels = ((width // sample_step) * (height // sample_step))
        red_ratio = red_pixel_count / sampled_pixels if sampled_pixels > 0 else 0
        
        is_red = red_ratio >= red_threshold
        
        return is_red
        
    except Exception:
        return False


def has_stamp_features(doc: Document, element, element_to_para=None, check_square=True, check_red=True, 
                       aspect_ratio_tolerance=0.03, red_threshold=0.3):
    """
    检查图片是否具有印章特征（近似方形或红色）
    
    Args:
        doc: Document对象
        element: 段落XML元素
        element_to_para: 元素到Paragraph对象的映射字典（可选）
        check_square: 是否检查是否为方形，默认True
        check_red: 是否检查是否为红色，默认True
        aspect_ratio_tolerance: 长宽比容差，默认0.03（3%容差）
        red_threshold: 红色像素占比阈值，默认0.3
        
    Returns:
        (是否具有印章特征, 特征详情字典) 元组
    """
    features = {
        'is_square': False,
        'is_red': False
    }
    
    if check_square:
        features['is_square'] = is_square_image(doc, element, element_to_para, aspect_ratio_tolerance)
    
    if check_red:
        features['is_red'] = is_red_image(doc, element, element_to_para, red_threshold)
    
    # 只要满足任一特征就认为是印章
    has_features = features['is_square'] or features['is_red']
    
    return (has_features, features)


def get_image_info(doc: Document, element, element_to_para=None):
    """
    获取图片的详细信息（用于debug模式）
    
    Args:
        doc: Document对象
        element: 段落XML元素
        element_to_para: 元素到Paragraph对象的映射字典（可选）
        
    Returns:
        包含图片详细信息的字典，如果无法获取信息则返回None
    """
    info = {
        'has_image_element': False,
        'image_type': None,
        'has_binary_data': False,
        'image_format': None,
        'content_type': None,
        'file_extension': None,
        'data_size': None,
        'image_size_emu': None,
        'image_size_pixels': None,
        'is_embedded': False,
        'is_linked': False,
        'relationship_id': None
    }
    
    # 检查是否有图片元素
    ns_drawing = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
    ns_main = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
    
    inline = element.find(f'.//{ns_drawing}inline')
    if inline is not None:
        info['has_image_element'] = True
        info['image_type'] = 'inline'
    else:
        blip = element.find(f'.//{ns_main}blip')
        if blip is not None:
            info['has_image_element'] = True
            info['image_type'] = 'blip'
    
    if not info['has_image_element']:
        return info
    
    # 找到对应的Paragraph对象
    if element_to_para is not None:
        para_obj = element_to_para.get(element)
    else:
        para_obj = None
        for para in doc.paragraphs:
            if para._element == element:
                para_obj = para
                break
    
    if para_obj is None:
        return info
    
    # 遍历段落中的所有runs，查找包含图片的run
    for run in para_obj.runs:
        inline = run._element.find(f'.//{ns_drawing}inline')
        if inline is not None:
            blip = inline.find(f'.//{ns_main}blip')
            if blip is not None:
                # 检查是否为嵌入图片
                r_embed = blip.get(qn('r:embed'))
                if r_embed:
                    info['is_embedded'] = True
                    info['relationship_id'] = r_embed
                    try:
                        # 获取图片数据
                        image_part = doc.part.related_parts[r_embed]
                        image_bytes = image_part.blob
                        info['has_binary_data'] = True
                        info['data_size'] = len(image_bytes)
                        info['content_type'] = image_part.content_type
                        
                        # 获取图片扩展名
                        ext_map = {
                            'image/png': '.png',
                            'image/jpeg': '.jpg',
                            'image/jpg': '.jpg',
                            'image/gif': '.gif',
                            'image/bmp': '.bmp',
                            'image/x-emf': '.emf',
                            'image/x-wmf': '.wmf',
                        }
                        info['file_extension'] = ext_map.get(image_part.content_type, '.png')
                        info['image_format'] = info['file_extension'].lstrip('.')
                        
                        # 尝试使用PIL获取像素尺寸
                        if PIL_AVAILABLE:
                            try:
                                img = Image.open(io.BytesIO(image_bytes))
                                info['image_size_pixels'] = f"{img.size[0]}x{img.size[1]}"
                            except Exception:
                                pass
                    except (KeyError, AttributeError, Exception) as e:
                        info['error'] = str(e)
                
                # 检查是否为链接图片
                r_link = blip.get(qn('r:link'))
                if r_link:
                    info['is_linked'] = True
                    info['relationship_id'] = r_link
                    info['has_binary_data'] = False
    
    # 获取EMU尺寸
    size_emu = get_image_size_from_paragraph(doc, element, element_to_para)
    if size_emu:
        info['image_size_emu'] = f"{size_emu[0]}x{size_emu[1]} EMU"
    
    return info
