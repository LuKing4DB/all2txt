import copy
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# 使用统一的日志配置
import sys
from pathlib import Path

# 添加src目录到路径
sys.path.insert(0, str(Path(__file__).parent.parent.parent))
from utils.logger import get_logger

logger = get_logger(__name__)


def find_mark_tags(text, pattern):
    """
    查找并处理字符串中所有<mark>...</mark>标签
    :param text: 输入字符串
    :param process_func: 自定义处理函数（默认为提取内容）
    :return: 处理后的结果列表或修改后的字符串
    """
    matches = pattern.finditer(text)

    results = []
    for match in matches:
        content = match.group(1)  # 提取标签内容
        results.append(content)  # 默认返回原始内容

    return results


def get_number_style_map(docx_path):
    """
    获取文档中所有自动编号的样式
    """
    docx = Document(docx_path)
    numbering_part = docx.part.numbering_part._element

    abstract_style = {}
    number2style = {}
    pattern_absNum = re.compile(r'<w:abstractNum (.*?)/w:abstractNum>', re.DOTALL)  # 非贪婪匹配，支持多行内容
    abstract_strs = find_mark_tags(str(numbering_part.xml), pattern_absNum)
    for abstract_str in abstract_strs:
        abs_id = re.search(r"w:abstractNumId=\"(\d+)\"", abstract_str)
        if not abs_id:
            continue
        pattern_lvl = re.compile(r'<w:lvl (.*?)/w:lvl>', re.DOTALL)  # 非贪婪匹配，支持多行内容
        lvl_strs = find_mark_tags(abstract_str, pattern_lvl)
        lvl_style = {}
        for lvl_str in lvl_strs:
            i_lvl = re.search(r"w:ilvl=\"(\d+)\"", lvl_str)
            if not i_lvl:
                continue
            i_start = re.search(r"w:start w:val=\"(.*)\"", lvl_str)
            i_numFmt = re.search(r"w:numFmt w:val=\"(.*)\"", lvl_str)
            i_lvlText = re.search(r"w:lvlText w:val=\"(.*)\"", lvl_str)
            i_lvlJc = re.search(r"w:lvlJc w:val=\"(.*)\"", lvl_str)
            # numId2style[(num_id, i_lvl.group(1))] = {
            lvl_style[i_lvl.group(1)] = {
                "i_start": i_start.group(1) if i_start else None,
                "i_numFmt": i_numFmt.group(1) if i_numFmt else None,
                "i_lvlText": i_lvlText.group(1) if i_lvlText else None,
                "i_lvlJc": i_lvlJc.group(1) if i_lvlJc else None,
            }
        abstract_style[abs_id.group(1)] = lvl_style

    pattern_num = re.compile(r'<w:num (.*?)/w:num>', re.DOTALL)
    number_strs = find_mark_tags(str(numbering_part.xml), pattern_num)
    for number_str in number_strs:
        num_id = re.search(r"w:numId=\"(\d+)\"", number_str)
        if not num_id:
            continue
        abs_id = re.search(r"w:abstractNumId w:val=\"(\d+)\"", number_str)
        if not abs_id:
            continue
        cur_style = copy.deepcopy(abstract_style.get(abs_id.group(1)))
        # 处理override
        pattern_override = re.compile(r'<w:lvlOverride (.*?)/w:lvlOverride>', re.DOTALL)
        override_strs = find_mark_tags(number_str, pattern_override)
        for override_str in override_strs:
            i_lvl = re.search(r"w:ilvl=\"(\d+)\"", override_str)
            if not i_lvl:
                continue
            cur_lvl_style = cur_style.get(i_lvl.group(1))
            i_start_override = re.search(r"w:startOverride w:val=\"(.*)\"", override_str)
            i_numFmt_override = re.search(r"w:numFmtOverride w:val=\"(.*)\"", override_str)
            i_lvlText_override = re.search(r"w:lvlTextOverride w:val=\"(.*)\"", override_str)
            i_lvlJc_override = re.search(r"w:lvlJcOverride w:val=\"(.*)\"", override_str)
            cur_style[i_lvl.group(1)] = {
                "i_start": i_start_override.group(1) if i_start_override else cur_lvl_style["i_start"],
                "i_numFmt": i_numFmt_override.group(1) if i_numFmt_override else cur_lvl_style["i_numFmt"],
                "i_lvlText": i_lvlText_override.group(1) if i_lvlText_override else cur_lvl_style["i_lvlText"],
                "i_lvlJc": i_lvlJc_override.group(1) if i_lvlJc_override else cur_lvl_style["i_lvlJc"],
            }
        number2style[num_id.group(1)] = cur_style

    # 整理输出表
    numId2style = {}
    for num_id, style in number2style.items():
        for lvl_id, lvl_style in style.items():
            numId2style[(num_id, lvl_id)] = lvl_style
    return numId2style


def format_number(number: int, num_fmt: str) -> str:
    """
    将数字转换为Word编号样式字符串
    :param number: 要转换的数字（>=1）
    :param num_fmt: 格式类型，如decimal/upperRoman等
    :return: 格式化后的字符串，无效格式返回空字符串
    """

    def _int_to_roman(n: int) -> str:
        val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
        syms = ["M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I"]
        roman = ""
        i = 0
        while n > 0:
            while n >= val[i]:
                roman += syms[i]
                n -= val[i]
            i += 1
        return roman

    def _int_to_alpha(n: int) -> str:
        letters = []
        while n > 0:
            n, rem = divmod(n - 1, 26)
            letters.append(chr(65 + rem))
        return ''.join(reversed(letters))

    def _int_to_chinese(n: int, upper: bool) -> str:
        # 数字和单位定义
        digits = ['零', '壹', '贰', '叁', '肆', '伍', '陆', '柒', '捌', '玖'] if upper else ['零', '一', '二', '三',
                                                                                             '四', '五', '六', '七',
                                                                                             '八', '九']
        units = ['', '拾', '佰', '仟'] if upper else ['', '十', '百', '千']
        big_units = ['', '万', '亿', '万亿']

        # 处理小节（每4位）
        sections = []
        while n > 0:
            sections.append(n % 10000)
            n = n // 10000

        result = []
        for i, sec in enumerate(reversed(sections)):
            section_str = []
            last_is_zero = False
            for j in range(4):
                num = sec % 10
                sec = sec // 10
                if num != 0:
                    section_str.append(digits[num] + units[j])
                    last_is_zero = False
                elif not last_is_zero and sec != 0:
                    section_str.append(digits[0])
                    last_is_zero = True
            # 添加节单位（万、亿等）
            if section_str:
                section_str = ''.join(reversed(section_str)).rstrip(digits[0])
                result.append(section_str + big_units[i])
        # 处理连续零和开头一十
        chinese = ''.join(reversed(result)).replace('零零', '零')
        return chinese if not chinese.startswith('一十') else chinese[1:]

    converters = {
        'decimal': lambda n: str(n),
        'upperRoman': lambda n: _int_to_roman(n),
        'lowerRoman': lambda n: _int_to_roman(n).lower(),
        'upperLetter': lambda n: _int_to_alpha(n),
        'lowerLetter': lambda n: _int_to_alpha(n).lower(),
        'chineseCounting': lambda n: _int_to_chinese(n, upper=False),
        'chineseCountingThousand': lambda n: _int_to_chinese(n, upper=False),
        'japaneseCounting': lambda n: _int_to_chinese(n, upper=False),
        'bullet': lambda n: '•'  # 标准项目符号
    }

    if number < 1:
        return ""

    return converters.get(num_fmt, lambda _: "")(number)


def find_percent_numbers(pattern_str: str) -> list:
    """
    提取字符串中所有 %数字 模式的数字

    :param pattern_str: 包含 %数字 模式的字符串，例如 "第%1章"、"%1.%2.%3"
    :return: 按出现顺序排列的数字列表，例如 [1,2,3]
    """
    return [int(match) for match in re.findall(r'%(\d+)', pattern_str)]


def style2prefix(style, offset):
    """
    将样式转换为标题前缀
    :param style: 样式字典
    :param offset: 偏离计数，用于表明当前标题序号
    """
    fmt_base = ["chineseCounting", "chineseCountingThousand", "japaneseCounting", "decimal", "lowerLetter", "lowerRoman", "upperLetter", "upperRoman", "bullet"]
    if not isinstance(style, dict):
        return ""
    start = style.get("i_start")
    num_fmt = style.get("i_numFmt")
    lvl_text = style.get("i_lvlText")
    # lvlJc = style.get("i_lvlJc")

    # 数据检查
    try:
        start = int(start)
        assert num_fmt in fmt_base
        assert "%" in lvl_text
    # except (ValueError, AssertionError):
    except:
        return ""

    # 根据样式生成前缀
    prefix = lvl_text
    pos2use = find_percent_numbers(lvl_text)
    max_pos = -1
    for pos in pos2use:
        max_pos = max(max_pos, pos)  # 找到最后一个用到的位置，方便增加偏移
    if max_pos > len(offset):  # 需要扩容
        offset += [0] * (max_pos - len(offset) + 1)
    num2change = {}
    for pos in pos2use:
        pattern = f"%{pos}"
        num2change[pattern] = format_number(offset[pos], num_fmt)
    max_pattern = f"%{max_pos}"
    num2change[max_pattern] = format_number(offset[max_pos] + start, num_fmt)  # 需要有初始限定
    for pattern, num in num2change.items():  # 替换所有占位符
        prefix = prefix.replace(pattern, num)
    offset[max_pos] += 1  # 增加计数
    return prefix + " "


def add_prefix_to_paragraph(ele, prefix):
    """在段落的最前面插入带编号前缀的run，处理分页符在开头的情况"""
    # 创建新run元素
    new_run = OxmlElement('w:r')

    # 复制原段落第一个run的格式（如果有的话）
    existing_runs = ele.xpath('.//w:r')
    if existing_runs:
        first_run_props = existing_runs[0].xpath('.//w:rPr')
        if first_run_props:
            new_run.append(copy.deepcopy(first_run_props[0]))

    # 创建文本元素
    new_text = OxmlElement('w:t')
    new_text.set(qn('xml:space'), 'preserve')  # 保留空格
    new_text.text = prefix
    new_run.append(new_text)

    # 智能定位插入位置
    pPr = ele.find(qn('w:pPr'))  # 查找段落属性

    if pPr is not None:
        # 查找pPr后的第一个run元素
        next_elem = pPr.getnext()
        while next_elem is not None:
            if next_elem.tag.endswith('r'):
                # 检查是否包含分页符
                if next_elem.find(qn('w:br')) is not None:
                    next_elem.addnext(new_run)  # 插入到分页符之后
                    return
                break
            next_elem = next_elem.getnext()

        # 默认插入到pPr之后
        pPr.addnext(new_run)
    else:
        # 检查第一个元素是否是分页符
        if len(ele) > 0 and ele[0].tag.endswith('r'):
            first_run = ele[0]
            if first_run.find(qn('w:br')) is not None:
                first_run.addnext(new_run)  # 插入到分页符之后
                return

        # 默认插入到段落开头
        ele.insert(0, new_run)


def remove_auto_numbering_from_styles(document):
    """
    从文档样式中移除自动编号属性
    :param document: python-docx的Document对象
    """
    # 遍历所有文档样式
    for style in document.styles:
        # 只处理段落样式
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue

        # 获取样式的XML元素
        style_element = style._element

        # 查找段落属性
        pPr = style_element.find(qn('w:pPr'))
        if pPr is not None:
            # 查找编号属性并移除
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)


def remove_auto_numbering(paragraph_element):
    """
    移除段落中的自动编号元素
    :param paragraph_element: docx.oxml.xmlchemy.OxmlElement对象
    """
    # 获取段落属性
    pPr = paragraph_element.pPr

    if pPr is None:
        return

    # 删除编号属性
    numPr = pPr.find(qn('w:numPr'))
    if numPr is not None:
        pPr.remove(numPr)

    # 删除隐藏的字段代码（如AUTONUM字段）
    for run in paragraph_element.xpath('.//w:r'):
        instr_text = run.find(qn('w:instrText'))
        if instr_text is not None and 'AUTONUM' in instr_text.text:
            # 删除整个字段代码结构
            parent = run.getparent()
            parent.remove(run)

            # 删除字段结果（可能存在的隐藏文本）
            next_sibling = parent.getnext()
            if next_sibling is not None and next_sibling.tag.endswith('r'):
                next_run = next_sibling
                if next_run.find(qn('w:lastRenderedPageBreak')) is None:
                    parent.getparent().remove(next_sibling)


def delete_auto_numbering_in_docx(docx_path):
    """
    删除docx文档中所有段落（包括表格内）的自动编号并替换为静态编号
    :param docx_path: docx文件路径
    :return: 处理后的Document对象
    """
    doc = Document(docx_path)
    try:
        id2style = get_number_style_map(docx_path)
    except NotImplementedError:
        logger.warning("未找到编号样式，直接返回原文档")
        return doc
    try:
        para_id2style = get_numbering_style_list(doc)
    except NotImplementedError:
        logger.warning("未找到编号样式，直接返回原文档")
    mark_json = {}

    # 递归处理所有段落元素的通用方法
    def process_paragraph(element):
        # 使用XPath查找所有段落元素（包括表格内的）
        for p in element.xpath('.//w:p'):
            if "sectPr" in str(p.xml):
                continue
            plan_b_flag = False
            if not hasattr(p, 'pPr') or p.pPr is None:
                plan_b_flag = True

            # 处理自动编号逻辑
            elif p.pPr.numPr is not None:
                test_str = str(p.xml)
                numId_re = re.search(r"<w:numId w:val=\"(\d+)\"", test_str)
                lvl_re = re.search(r"<w:ilvl w:val=\"(\d+)\"", test_str)

                if numId_re and lvl_re:
                    numId = numId_re.group(1)
                    lvl = lvl_re.group(1)

                    if numId == "0":
                        continue
                    else:

                        couple = (numId, lvl)
                        if couple not in mark_json:
                            mark_json[couple] = [0] * 20
                        offset = mark_json[couple]

                        style = id2style.get(couple)
                        prefix = style2prefix(style, offset)
                        if prefix:
                            add_prefix_to_paragraph(p, prefix)
                            remove_auto_numbering(p)
                        else:
                            logger.warning("警告：未生成编号前缀，自动编号未被移除，避免误删内容")
                        # if prefix:
                        #     add_prefix_to_paragraph(p, prefix)

                        # remove_auto_numbering(p)

                        # 更新计数器
                        if couple in mark_json:
                            mark_json[couple][int(lvl)] += 1
                        # 清除对应子层级计数器
                        for i in range(int(lvl) + 1, 20):
                            cur_couple = (numId, str(i))
                            if cur_couple in mark_json:
                                del mark_json[cur_couple]

            else:
                plan_b_flag = True

            if plan_b_flag:

                style_name = p.style
                if isinstance(style_name, str):
                    couple = para_id2style.get(style_name)
                    if couple is not None:
                        numId = couple[0]
                        lvl = couple[1]
                        if numId == "0":
                            continue
                        if couple not in mark_json:
                            mark_json[couple] = [0] * 20
                        offset = mark_json[couple]
                        style = id2style.get(couple)
                        prefix = style2prefix(style, offset)
                        if prefix:
                            add_prefix_to_paragraph(p, prefix)

                        # remove_auto_numbering(p)
                        # 更新计数器
                        if couple in mark_json:
                            mark_json[couple][int(lvl)] += 1
                        # 清除对应子层级计数器
                        for i in range(int(lvl) + 1, 20):
                            cur_couple = (numId, str(i))
                            if cur_couple in mark_json:
                                del mark_json[cur_couple]

        # 删去原文自动编号属性
        remove_auto_numbering_from_styles(doc)

    # 处理文档主体和表格
    process_paragraph(doc.element.body)

    # 特别处理表格中的嵌套段落
    # for table in doc.tables:
    #     try:
    #         for row in table.rows:
    #             for cell in row.cells:
    #                 # 递归处理单元格内容
    #                 process_paragraph(cell._element)
    #     except:
    #         continue

    return doc


def get_numbering_style_list(docx):
    numbering_style_list = {}
    try:
        for style in docx.styles.element.style_lst:
            style_id = style.styleId
            if style_id is not None or style_id != "":
                cur_pPr = style.pPr
                if cur_pPr is not None:
                    numPr = cur_pPr.numPr
                    if numPr is not None:
                        soup_text = str(numPr.xml)
                        numId_re = re.search(r"<w:numId w:val=\"(\d+)\"", soup_text)
                        lvl_re = re.search(r"<w:ilvl w:val=\"(\d+)\"", soup_text)
                        if numId_re and lvl_re:
                            numId = numId_re.group(1)
                            lvl = lvl_re.group(1)
                            if numId == "0":
                                continue
                            couple = (numId, lvl)
                            numbering_style_list[style_id] = couple
                        elif numId_re:
                            numId = numId_re.group(1)
                            lvl = "0"
                            if numId == "0":
                                continue
                            couple = (numId, lvl)
                            numbering_style_list[style_id] = couple
    except Exception as e:
        logger.warning(e)

    return numbering_style_list


def run_delete_auto_numbering_in_docx(docx_path: str) -> None:
    try:
        file_doc_obj = delete_auto_numbering_in_docx(docx_path)
        file_doc_obj.save(docx_path)
    except Exception as e:
        logger.error(str(e))


if __name__ == '__main__':
    test_docx_path = r"tender.docx"
    test_docx = delete_auto_numbering_in_docx(test_docx_path)
    test_docx.save(r"D:\files\out\img_disappear_out.docx")
